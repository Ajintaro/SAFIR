#!/usr/bin/env python3
"""
CGI AFCEA San-Feldeinsatz — Web-Dashboard
FastAPI Backend mit WebSocket, Templates und Vosk-Sprachsteuerung.
"""

import asyncio
import json
import os
import subprocess
import tempfile
import time
import threading
from datetime import datetime
from pathlib import Path

import numpy as np
import psutil
import sounddevice as sd
import soundfile as sf
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.requests import Request

import signal
import httpx

PROJECT_DIR = Path(__file__).parent
WHISPER_CLI = PROJECT_DIR / "whisper.cpp" / "build" / "bin" / "whisper-cli"
WHISPER_SERVER = PROJECT_DIR / "whisper.cpp" / "build" / "bin" / "whisper-server"
MODELS_DIR = PROJECT_DIR / "models"
WHISPER_SERVER_PORT = 8178  # interner Port fuer whisper-server
PROTOCOLS_DIR = PROJECT_DIR / "protocols"
PROTOCOLS_DIR.mkdir(exist_ok=True)
VOSK_MODEL_PATH = MODELS_DIR / "vosk-model-small-de"
OLLAMA_URL = "http://127.0.0.1:11434"
OLLAMA_MODEL = "qwen2.5:1.5b"

SAMPLE_RATE = 16000

app = FastAPI(title="CGI San-Feldeinsatz")
app.mount("/static", StaticFiles(directory=str(PROJECT_DIR / "static")), name="static")
templates = Jinja2Templates(directory=str(PROJECT_DIR / "templates"))


# ---------------------------------------------------------------------------
# Templates fuer Patientenakten
# ---------------------------------------------------------------------------
RECORD_TEMPLATES = {
    "tccc": {
        "id": "tccc",
        "name": "TCCC Card",
        "description": "Tactical Combat Casualty Care — Verwundetenkarte",
        "icon": "&#9764;",
        "sections": [
            {
                "title": "Triage",
                "fields": [
                    {"key": "triage_cat", "label": "Triage-Kategorie", "type": "select",
                     "options": ["T1 — Sofort (Rot)", "T2 — Aufgeschoben (Gelb)", "T3 — Leicht (Gruen)", "T4 — Abwartend (Blau)"]},
                    {"key": "triage_time", "label": "Zeitpunkt Sichtung", "type": "text", "default": ""},
                    {"key": "evac_priority", "label": "Transportprioritaet", "type": "select",
                     "options": ["Urgent", "Urgent Surgical", "Priority", "Routine"]},
                ],
            },
            {
                "title": "Verletzung",
                "fields": [
                    {"key": "mechanism", "label": "Mechanismus", "type": "multiselect",
                     "options": ["Schuss (GSW)", "IED/Explosion", "Splitter", "Verbrennung", "Sturz", "Stumpfes Trauma", "Mine", "Granate", "Fahrzeugunfall", "Sonstiges"]},
                    {"key": "injury_location", "label": "Verletzungsort", "type": "text", "default": ""},
                    {"key": "injury_type", "label": "Verletzungsart", "type": "multiselect",
                     "options": ["Verwundung", "Verbrennung", "Erkrankung", "Vergiftung", "Strahlung", "Psych. Belastung"]},
                ],
            },
            {
                "title": "Vitalzeichen",
                "fields": [
                    {"key": "pulse", "label": "Puls (bpm)", "type": "text", "default": ""},
                    {"key": "bp", "label": "Blutdruck (sys/dia)", "type": "text", "default": ""},
                    {"key": "resp_rate", "label": "Atemfrequenz (/min)", "type": "text", "default": ""},
                    {"key": "spo2", "label": "SpO2 (%)", "type": "text", "default": ""},
                    {"key": "avpu", "label": "Bewusstsein (AVPU)", "type": "select",
                     "options": ["A — Wach", "V — Reagiert auf Ansprache", "P — Reagiert auf Schmerz", "U — Bewusstlos"]},
                    {"key": "pain", "label": "Schmerz (0-10)", "type": "text", "default": ""},
                ],
            },
            {
                "title": "Behandlung",
                "fields": [
                    {"key": "tourniquet", "label": "Tourniquet", "type": "text", "default": ""},
                    {"key": "tourniquet_time", "label": "TQ-Zeit", "type": "text", "default": ""},
                    {"key": "hemostatic", "label": "Haemostatikum / Verband", "type": "text", "default": ""},
                    {"key": "airway", "label": "Atemwegssicherung", "type": "select",
                     "options": ["Keine", "NPA", "SGA", "Endotracheal", "Koniotomie"]},
                    {"key": "chest_seal", "label": "Chest Seal / Thorax", "type": "text", "default": ""},
                    {"key": "iv_io", "label": "Zugang (IV/IO)", "type": "text", "default": ""},
                    {"key": "fluids", "label": "Infusionen", "type": "text", "default": ""},
                    {"key": "medications", "label": "Medikamente", "type": "textarea", "default": ""},
                    {"key": "splint", "label": "Schienung", "type": "text", "default": ""},
                ],
            },
        ],
    },
    "9liner": {
        "id": "9liner",
        "name": "9-Liner MEDEVAC",
        "description": "NATO MEDEVAC-Anforderung",
        "icon": "&#9992;",
        "sections": [
            {
                "title": "MEDEVAC Request",
                "fields": [
                    {"key": "line1", "label": "Line 1 — Koordinaten Landezone", "type": "text", "default": ""},
                    {"key": "line2", "label": "Line 2 — Funkfrequenz / Rufzeichen", "type": "text", "default": ""},
                    {"key": "line3", "label": "Line 3 — Patienten nach Dringlichkeit", "type": "text", "default": ""},
                    {"key": "line4", "label": "Line 4 — Sonderausstattung", "type": "select",
                     "options": ["A — Keine", "B — Winde (Hoist)", "C — Bergungsgeraet", "D — Beatmungsgeraet"]},
                    {"key": "line5", "label": "Line 5 — Patienten (Liegend/Gehfaehig)", "type": "text", "default": ""},
                    {"key": "line6", "label": "Line 6 — Sicherheitslage", "type": "select",
                     "options": ["N — Kein Feind", "P — Moegl. Feind", "E — Feind im Gebiet", "X — Bewaffnete Eskorte"]},
                    {"key": "line7", "label": "Line 7 — Markierung Landeplatz", "type": "select",
                     "options": ["A — Panels", "B — Pyrotechnik", "C — Rauch", "D — Keine", "E — Sonstige"]},
                    {"key": "line8", "label": "Line 8 — Nationalitaet / Status", "type": "text", "default": ""},
                    {"key": "line9", "label": "Line 9 — ABC / Gelaende", "type": "text", "default": ""},
                ],
            },
        ],
    },
    "erstbefund": {
        "id": "erstbefund",
        "name": "Erstbefund",
        "description": "Erste Untersuchung — Anamnese & Befund",
        "icon": "&#9879;",
        "sections": [
            {
                "title": "Stammdaten",
                "fields": [
                    {"key": "rank", "label": "Dienstgrad", "type": "text", "default": ""},
                    {"key": "unit", "label": "Einheit", "type": "text", "default": ""},
                    {"key": "dob", "label": "Geburtsdatum", "type": "text", "default": ""},
                    {"key": "blood_type", "label": "Blutgruppe", "type": "select",
                     "options": ["A+", "A-", "B+", "B-", "AB+", "AB-", "0+", "0-", "Unbekannt"]},
                    {"key": "allergies", "label": "Allergien / Unvertraeglichkeiten", "type": "text", "default": ""},
                ],
            },
            {
                "title": "Anamnese",
                "fields": [
                    {"key": "complaint", "label": "Hauptbeschwerde", "type": "textarea", "default": ""},
                    {"key": "history", "label": "Vorgeschichte / Vorerkrankungen", "type": "textarea", "default": ""},
                    {"key": "current_meds", "label": "Aktuelle Medikation", "type": "text", "default": ""},
                ],
            },
            {
                "title": "Befund",
                "fields": [
                    {"key": "general", "label": "Allgemeinzustand", "type": "textarea", "default": ""},
                    {"key": "pulse", "label": "Puls (bpm)", "type": "text", "default": ""},
                    {"key": "bp", "label": "Blutdruck", "type": "text", "default": ""},
                    {"key": "resp_rate", "label": "Atemfrequenz", "type": "text", "default": ""},
                    {"key": "spo2", "label": "SpO2 (%)", "type": "text", "default": ""},
                    {"key": "temp", "label": "Temperatur", "type": "text", "default": ""},
                ],
            },
            {
                "title": "Diagnose & Therapie",
                "fields": [
                    {"key": "diagnosis", "label": "Diagnose", "type": "textarea", "default": ""},
                    {"key": "therapy", "label": "Therapie / Massnahmen", "type": "textarea", "default": ""},
                    {"key": "disposition", "label": "Weiteres Vorgehen", "type": "textarea", "default": ""},
                ],
            },
        ],
    },
    "mist": {
        "id": "mist",
        "name": "MIST Uebergabe",
        "description": "Standardisierte Patientenuebergabe",
        "icon": "&#8644;",
        "sections": [
            {
                "title": "MIST Schema",
                "fields": [
                    {"key": "m_mechanism", "label": "M — Mechanismus (Wie ist es passiert?)", "type": "textarea", "default": ""},
                    {"key": "i_injuries", "label": "I — Injuries (Welche Verletzungen?)", "type": "textarea", "default": ""},
                    {"key": "s_signs", "label": "S — Signs/Symptoms (Vitalzeichen, Bewusstsein)", "type": "textarea", "default": ""},
                    {"key": "t_treatment", "label": "T — Treatment (Was wurde gemacht?)", "type": "textarea", "default": ""},
                ],
            },
            {
                "title": "Uebergabe-Details",
                "fields": [
                    {"key": "from_unit", "label": "Uebergebende Einheit", "type": "text", "default": ""},
                    {"key": "to_unit", "label": "Aufnehmende Einrichtung", "type": "text", "default": ""},
                    {"key": "transport", "label": "Transportmittel", "type": "text", "default": ""},
                    {"key": "handover_time", "label": "Zeitpunkt Uebergabe", "type": "text", "default": ""},
                ],
            },
        ],
    },
    "freitext": {
        "id": "freitext",
        "name": "Freitext",
        "description": "Freie Spracheingabe ohne Struktur",
        "icon": "&#9998;",
        "sections": [],
    },
}


# ---------------------------------------------------------------------------
# Vosk Voice Commands
# ---------------------------------------------------------------------------
VOICE_COMMANDS = {
    "aufnahme starten": "record_start",
    "aufnahme start": "record_start",
    "start aufnahme": "record_start",
    "starte aufnahme": "record_start",
    "aufnahme stoppen": "record_stop",
    "aufnahme stopp": "record_stop",
    "aufnahme stop": "record_stop",
    "stopp aufnahme": "record_stop",
    "stop aufnahme": "record_stop",
    "stoppe aufnahme": "record_stop",
    "neue session": "new_session",
    "neuer patient": "new_session",
    "neue akte": "new_session",
    "export": "export_docx",
    "exportieren": "export_docx",
    "dokument erstellen": "export_docx",
    "akte exportieren": "export_docx",
}


# ---------------------------------------------------------------------------
# State
# ---------------------------------------------------------------------------
class AppState:
    def __init__(self):
        self.current_model = None
        self.model_path = None
        self.model_loaded = False
        self.model_loading = False
        self.whisper_process = None
        self.recording = False
        self.transcribing = False
        self.audio_device = None
        self.audio_chunks = []
        self.stream = None
        self.sessions = {}
        self.active_session = None
        self.ws_clients: list[WebSocket] = []
        self.language = "de"
        self.model_ram_mb = 0
        # Vosk
        self.vosk_enabled = False
        self.vosk_model = None
        self.vosk_recognizer = None
        self.vosk_listening = False
        self.persistent_stream = None
        self.event_loop = None
        self.vosk_command_queue = []  # thread-safe command queue

    def available_models(self):
        models = []
        for f in sorted(MODELS_DIR.glob("ggml-*.bin")):
            size_mb = f.stat().st_size / (1024 * 1024)
            name = f.stem.replace("ggml-", "")
            models.append({"name": name, "file": f.name, "size_mb": round(size_mb)})
        return models

    def audio_devices(self):
        devices = []
        for i, dev in enumerate(sd.query_devices()):
            if dev["max_input_channels"] > 0:
                devices.append({
                    "id": i,
                    "name": dev["name"],
                    "channels": dev["max_input_channels"],
                    "default": i == sd.default.device[0],
                })
        return devices


state = AppState()


async def broadcast(msg: dict):
    """Sendet JSON an alle verbundenen WebSocket-Clients."""
    dead = []
    for ws in state.ws_clients:
        try:
            await ws.send_json(msg)
        except Exception:
            dead.append(ws)
    for ws in dead:
        state.ws_clients.remove(ws)


# ---------------------------------------------------------------------------
# System Monitoring
# ---------------------------------------------------------------------------
def get_system_stats():
    cpu_percent = psutil.cpu_percent(interval=0)
    cpu_freq = psutil.cpu_freq()
    mem = psutil.virtual_memory()

    gpu_usage = "N/A"
    gpu_mem_used = 0
    gpu_mem_total = 0

    try:
        gpu_load_path = Path("/sys/devices/platform/bus@0/17000000.gpu/load")
        if gpu_load_path.exists():
            raw = int(gpu_load_path.read_text().strip())
            gpu_usage = str(round(raw / 10, 1))
            gpu_mem_used = round(mem.used / 1024 / 1024)
            gpu_mem_total = round(mem.total / 1024 / 1024)
    except Exception:
        pass

    temps = {}
    try:
        for tz in Path("/sys/class/thermal/").glob("thermal_zone*"):
            try:
                name = (tz / "type").read_text().strip()
                temp = int((tz / "temp").read_text().strip()) / 1000
                temps[name] = round(temp, 1)
            except Exception:
                pass
    except Exception:
        pass

    disk = psutil.disk_usage("/")

    return {
        "cpu_percent": cpu_percent,
        "cpu_freq_mhz": round(cpu_freq.current) if cpu_freq else 0,
        "cpu_cores": psutil.cpu_count(),
        "ram_used_mb": round(mem.used / 1024 / 1024),
        "ram_total_mb": round(mem.total / 1024 / 1024),
        "ram_percent": mem.percent,
        "gpu_usage": gpu_usage,
        "gpu_mem_used_mb": gpu_mem_used,
        "gpu_mem_total_mb": gpu_mem_total,
        "temperatures": temps,
        "disk_used_gb": round(disk.used / 1024**3, 1),
        "disk_total_gb": round(disk.total / 1024**3, 1),
        "disk_percent": round(disk.percent, 1),
    }


# ---------------------------------------------------------------------------
# Vosk Keyword Detection
# ---------------------------------------------------------------------------
def init_vosk():
    """Initialisiert Vosk fuer Sprachbefehle."""
    if not VOSK_MODEL_PATH.exists():
        print("Vosk-Modell nicht gefunden, Sprachbefehle deaktiviert")
        return False
    try:
        from vosk import Model, KaldiRecognizer, SetLogLevel
        SetLogLevel(-1)
        state.vosk_model = Model(str(VOSK_MODEL_PATH))
        state.vosk_recognizer = KaldiRecognizer(state.vosk_model, SAMPLE_RATE)
        state.vosk_enabled = True
        print(f"Vosk Sprachsteuerung aktiviert (Modell: {VOSK_MODEL_PATH.name})")
        return True
    except Exception as e:
        print(f"Vosk Init fehlgeschlagen: {e}")
        return False


def match_voice_command(text: str) -> str | None:
    """Matched erkannten Text gegen Sprachbefehle. Gibt Aktions-ID zurueck."""
    text = text.lower().strip()
    if len(text) < 3:
        return None
    # Exakter Match
    if text in VOICE_COMMANDS:
        return VOICE_COMMANDS[text]
    # Substring-Match
    for phrase, action in VOICE_COMMANDS.items():
        if phrase in text or text in phrase:
            return action
    return None


def persistent_audio_callback(indata, frames, time_info, status):
    """Audio-Callback fuer persistenten Stream — fuettert Vosk und Recording."""
    # Recording-Buffer fuellen
    if state.recording:
        state.audio_chunks.append(indata.copy())

    # Vosk fuettern — immer wenn enabled (auch waehrend Aufnahme fuer Stop-Befehl)
    if state.vosk_enabled and state.vosk_recognizer and not state.transcribing:
        try:
            data = (indata[:, 0] * 32767).astype(np.int16).tobytes()
            if state.vosk_recognizer.AcceptWaveform(data):
                result = json.loads(state.vosk_recognizer.Result())
                text = result.get("text", "").strip()
                if text:
                    action = match_voice_command(text)
                    if action:
                        # Waehrend Aufnahme nur Stop-Befehl erlauben
                        if state.recording and action != "record_stop":
                            return
                        # Wenn nicht aufnimmt, kein Stop-Befehl
                        if not state.recording and action == "record_stop":
                            return
                        state.vosk_command_queue.append({
                            "action": action,
                            "text": text,
                            "time": datetime.now().isoformat(),
                        })
        except Exception:
            pass


def start_persistent_stream():
    """Startet den persistenten Audio-Stream."""
    stop_persistent_stream()
    try:
        state.persistent_stream = sd.InputStream(
            samplerate=SAMPLE_RATE, channels=1, dtype="float32",
            blocksize=int(SAMPLE_RATE * 0.1),
            device=state.audio_device, callback=persistent_audio_callback,
        )
        state.persistent_stream.start()
        state.vosk_listening = state.vosk_enabled
        print(f"Persistenter Audio-Stream gestartet (Device: {state.audio_device or 'default'})")
        return True
    except Exception as e:
        print(f"Audio-Stream Fehler: {e}")
        return False


def stop_persistent_stream():
    """Stoppt den persistenten Audio-Stream."""
    if state.persistent_stream:
        try:
            state.persistent_stream.stop()
            state.persistent_stream.close()
        except Exception:
            pass
        state.persistent_stream = None
    state.vosk_listening = False


async def process_vosk_commands():
    """Verarbeitet Vosk-Sprachbefehle aus der Queue (laeuft als asyncio task)."""
    while True:
        if state.vosk_command_queue:
            cmd = state.vosk_command_queue.pop(0)
            action = cmd["action"]
            text = cmd["text"]
            print(f"Sprachbefehl: '{text}' -> {action}")

            await broadcast({
                "type": "voice_command",
                "action": action,
                "text": text,
            })

            try:
                if action == "record_start":
                    if not state.recording and not state.transcribing:
                        # 1.5s warten damit der Sprachbefehl selbst nicht aufgenommen wird
                        await asyncio.sleep(1.5)
                        state.audio_chunks = []  # Buffer leeren
                        await start_recording_internal()
                elif action == "record_stop":
                    if state.recording:
                        # Letzte 1.5s Audio abschneiden (enthaelt "stopp"-Befehl)
                        trim_chunks = int(1.5 / 0.1)  # 1.5s bei 100ms blocksize
                        if len(state.audio_chunks) > trim_chunks:
                            state.audio_chunks = state.audio_chunks[:-trim_chunks]
                        await stop_recording()
                elif action == "new_session":
                    await create_session_internal("Unbekannt", "Feld", "", "freitext")
                elif action == "export_docx":
                    pass  # Export braucht Download — nur UI-Hinweis
            except Exception as e:
                print(f"Vosk Befehl Fehler: {e}")

        await asyncio.sleep(0.1)


# ---------------------------------------------------------------------------
# Whisper Server Lifecycle
# ---------------------------------------------------------------------------
def start_whisper_server(model_path: Path) -> bool:
    """Startet den whisper-server mit dem gegebenen Modell."""
    stop_whisper_server()

    env = os.environ.copy()
    env["PATH"] = "/usr/local/cuda/bin:" + env.get("PATH", "")

    cmd = [
        str(WHISPER_SERVER),
        "-m", str(model_path),
        "-l", state.language,
        "-t", "4",
        "-nfa",
        "-bs", "1",
        "-bo", "1",
        "--host", "127.0.0.1",
        "--port", str(WHISPER_SERVER_PORT),
        "--convert",
    ]

    try:
        state.whisper_process = subprocess.Popen(
            cmd, env=env,
            stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        )
        for _ in range(30):
            time.sleep(1)
            if state.whisper_process.poll() is not None:
                stderr = state.whisper_process.stderr.read().decode()
                stderr = "\n".join(l for l in stderr.splitlines() if "NvMap" not in l)
                print(f"whisper-server Fehler: {stderr[:300]}")
                return False
            try:
                r = httpx.get(f"http://127.0.0.1:{WHISPER_SERVER_PORT}/", timeout=2)
                if r.status_code == 200:
                    state.model_loaded = True
                    try:
                        proc = psutil.Process(state.whisper_process.pid)
                        state.model_ram_mb = round(proc.memory_info().rss / 1024 / 1024)
                    except Exception:
                        state.model_ram_mb = round(model_path.stat().st_size / 1024 / 1024)
                    print(f"whisper-server bereit (PID {state.whisper_process.pid}, ~{state.model_ram_mb} MB)")
                    return True
            except Exception:
                continue
        return False
    except Exception as e:
        print(f"whisper-server Start fehlgeschlagen: {e}")
        return False


def stop_whisper_server():
    """Stoppt den whisper-server und gibt GPU-Speicher frei."""
    if state.whisper_process:
        try:
            state.whisper_process.terminate()
            state.whisper_process.wait(timeout=5)
        except Exception:
            state.whisper_process.kill()
        state.whisper_process = None
    state.model_loaded = False
    state.model_ram_mb = 0


def is_whisper_server_alive() -> bool:
    """Prueft ob der whisper-server laeuft."""
    if not state.whisper_process or state.whisper_process.poll() is not None:
        state.model_loaded = False
        return False
    try:
        r = httpx.get(f"http://127.0.0.1:{WHISPER_SERVER_PORT}/", timeout=2)
        return r.status_code == 200
    except Exception:
        return False


# ---------------------------------------------------------------------------
# Transcription
# ---------------------------------------------------------------------------
def run_transcribe(audio: np.ndarray, language: str = "de") -> dict:
    """Transkribiert Audio via whisper-server HTTP API."""
    if not state.model_loaded:
        return {"error": "Kein Modell geladen", "text": "", "duration": 0}

    if audio.ndim > 1:
        audio = audio[:, 0]
    peak = np.max(np.abs(audio))
    if peak > 0 and peak < 0.01:
        audio = audio * (0.5 / peak)

    with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
        wav_path = f.name
        sf.write(wav_path, audio, SAMPLE_RATE, subtype='PCM_16')

    try:
        audio_duration = len(audio) / SAMPLE_RATE
        start = time.time()

        with open(wav_path, "rb") as f:
            response = httpx.post(
                f"http://127.0.0.1:{WHISPER_SERVER_PORT}/inference",
                files={"file": ("audio.wav", f, "audio/wav")},
                data={
                    "response_format": "json",
                    "language": language,
                },
                timeout=120,
            )

        elapsed = time.time() - start

        if response.status_code != 200:
            return {"error": f"Server-Fehler: {response.status_code}", "text": "", "duration": 0}

        result = response.json()
        text = result.get("text", "").strip()

        if not text:
            text = "(Stille / nicht erkannt)"

        rtf = elapsed / audio_duration if audio_duration > 0 else 0

        return {
            "text": text,
            "processing_time": round(elapsed, 2),
            "audio_duration": round(audio_duration, 2),
            "rtf": round(rtf, 3),
        }
    except httpx.TimeoutException:
        return {"error": "Transkription Timeout", "text": "", "duration": 0}
    except Exception as e:
        return {"error": str(e)[:300], "text": "", "duration": 0}
    finally:
        os.unlink(wav_path)


# ---------------------------------------------------------------------------
# LLM Field Extraction (Ollama / Qwen)
# ---------------------------------------------------------------------------
def build_extraction_prompt(template_id: str, text: str) -> str:
    """Baut den Prompt fuer die LLM-Feldextraktion mit Few-Shot Beispiel."""
    tpl = RECORD_TEMPLATES.get(template_id)
    if not tpl or not tpl.get("sections"):
        return ""

    # Sammle alle Felder mit Keys und Labels
    fields_desc = []
    field_keys = []
    for section in tpl["sections"]:
        for field in section["fields"]:
            key = field["key"]
            label = field["label"]
            field_keys.append(key)
            if field["type"] == "select":
                opts = ", ".join(field["options"])
                fields_desc.append(f"- {key}: {label} (Optionen: {opts})")
            else:
                fields_desc.append(f"- {key}: {label}")

    fields_list = "\n".join(fields_desc)

    # Few-Shot Beispiel fuer 9-Liner
    example = ""
    if template_id == "9liner":
        example = """
Beispiel:
Text: Standort Grid 12345678. Rufzeichen Alpha1 auf 45.5 MHz. Ein Verwundeter dringend. Brauche Trage. Ein liegender Patient. Kein Feind. Rauchzeichen. NATO-Soldat. Offenes Gelaende, keine Kontaminierung.
JSON: {"line1": "Grid 12345678", "line2": "Alpha1, 45.5 MHz", "line3": "1 dringend", "line4": "A — Keine", "line5": "1 liegend", "line6": "N — Kein Feind", "line7": "C — Rauch", "line8": "NATO-Soldat", "line9": "Offenes Gelaende, keine Kontaminierung"}
"""

    return f"""Du bist ein militaerischer Sanitaets-Assistent. Extrahiere aus dem Text die Felder als JSON.

Felder:
{fields_list}

Regeln:
- Nutze NUR Informationen aus dem Text
- Felder ohne Info: leerer String ""
- Kurze, praezise Werte
- Bei Select-Feldern: nutze die passende Option
{example}
Text: {text}

JSON:"""


def run_llm_extraction(template_id: str, text: str) -> dict:
    """Extrahiert Template-Felder aus Text via Ollama LLM (CPU-only)."""
    prompt = build_extraction_prompt(template_id, text)
    if not prompt:
        return {}

    try:
        response = httpx.post(
            f"{OLLAMA_URL}/api/generate",
            json={
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
                "format": "json",
                "options": {"num_gpu": 0, "temperature": 0.1, "num_predict": 500},
            },
            timeout=120,
        )
        if response.status_code == 200:
            result = response.json()
            raw = result.get("response", "{}")
            try:
                extracted = json.loads(raw)
                print(f"LLM Extraktion: {len(extracted)} Felder extrahiert")
                return extracted
            except json.JSONDecodeError:
                print(f"LLM JSON Parse Fehler: {raw[:200]}")
                return {}
        else:
            print(f"LLM Fehler: {response.status_code}")
            return {}
    except Exception as e:
        print(f"LLM Extraktion fehlgeschlagen: {e}")
        return {}


# ---------------------------------------------------------------------------
# Document Generation
# ---------------------------------------------------------------------------
def _docx_add_header(doc, title):
    """Einheitlicher DOCX-Header."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph()


def _docx_add_patient_table(doc, session):
    """Patienten-Stammdaten Tabelle."""
    tpl = RECORD_TEMPLATES.get(session.get("template_id", "freitext"), {})
    base_fields = [
        ("Patient", session.get("patient_name", "Unbekannt")),
        ("Datum", session.get("date", datetime.now().strftime("%Y-%m-%d %H:%M"))),
        ("Einsatzort", session.get("location", "Feld")),
        ("Sanitaeter", session.get("medic", "—")),
        ("Dokumenttyp", tpl.get("name", "Freitext")),
    ]
    table = doc.add_table(rows=len(base_fields), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(base_fields):
        cell0 = table.rows[i].cells[0]
        cell0.text = label
        for p in cell0.paragraphs:
            for r in p.runs:
                r.bold = True
        table.rows[i].cells[1].text = str(value)
    doc.add_paragraph()


def _docx_add_template_data(doc, session):
    """Template-spezifische Felder ins DOCX schreiben."""
    template_id = session.get("template_id", "freitext")
    tpl = RECORD_TEMPLATES.get(template_id)
    tdata = session.get("template_data", {})

    if not tpl or not tpl.get("sections") or not tdata:
        return

    for section in tpl["sections"]:
        doc.add_heading(section["title"], level=2)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for field in section["fields"]:
            val = tdata.get(field["key"], "")
            if isinstance(val, list):
                val = ", ".join(val)
            if val:
                row = table.add_row()
                cell0 = row.cells[0]
                cell0.text = field["label"]
                for p in cell0.paragraphs:
                    for r in p.runs:
                        r.bold = True
                row.cells[1].text = str(val)
        if len(table.rows) == 0:
            # Leere Section — Platzhalter
            row = table.add_row()
            row.cells[0].text = "(Keine Daten)"
        doc.add_paragraph()


def _docx_add_transcripts(doc, session):
    """Transkriptions-Eintraege ins DOCX."""
    records = session.get("records", [])
    if records:
        doc.add_heading("Spracheingabe / Transkription", level=2)
        for i, record in enumerate(records, 1):
            para = doc.add_paragraph()
            run_time = para.add_run(f"[{record['time']}] ")
            run_time.bold = True
            run_time.font.size = Pt(9)
            para.add_run(record["text"])
        doc.add_paragraph()


def _docx_add_footer(doc):
    """DOCX Footer."""
    doc.add_paragraph()
    doc.add_heading("Unterschrift Sanitaeter", level=2)
    doc.add_paragraph()
    doc.add_paragraph("_" * 30)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("CGI Deutschland — AFCEA Demonstrator — KI-gestuetzte Felddokumentation (SAFIR)")
    run.font.size = Pt(8)
    run.italic = True


def generate_docx(session: dict) -> Path:
    """Generiert eine Patienten-Akte als Word-Dokument."""
    doc = Document()
    template_id = session.get("template_id", "freitext")
    tpl = RECORD_TEMPLATES.get(template_id, RECORD_TEMPLATES["freitext"])

    # Triage-Farbe in Header fuer TCCC
    tdata = session.get("template_data", {})
    triage = tdata.get("triage_cat", "")

    if template_id == "tccc":
        title = "TCCC CASUALTY CARD"
        if "T1" in triage:
            title += " — SOFORT (T1)"
        elif "T2" in triage:
            title += " — AUFGESCHOBEN (T2)"
        elif "T3" in triage:
            title += " — LEICHT (T3)"
        elif "T4" in triage:
            title += " — ABWARTEND (T4)"
    elif template_id == "9liner":
        title = "9-LINER MEDEVAC ANFORDERUNG"
    elif template_id == "mist":
        title = "MIST PATIENTENUEBERGABE"
    elif template_id == "erstbefund":
        title = "ERSTBEFUND — SANITAETSDIENST"
    else:
        title = "SANITAETSDIENST — PATIENTENAKTE"

    _docx_add_header(doc, title)
    _docx_add_patient_table(doc, session)
    _docx_add_template_data(doc, session)
    _docx_add_transcripts(doc, session)
    _docx_add_footer(doc)

    patient = session.get("patient_name", "patient").replace(" ", "_")
    filename = f"{template_id}_{patient}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = PROTOCOLS_DIR / filename
    doc.save(str(filepath))
    return filepath


# ---------------------------------------------------------------------------
# Internal helpers (used by both API and Vosk)
# ---------------------------------------------------------------------------
async def start_recording_internal():
    """Startet Aufnahme (intern, ohne HTTP)."""
    if state.recording or state.transcribing:
        return
    if not state.model_path:
        return

    state.audio_chunks = []
    state.recording = True
    # Vosk pausiert waehrend Aufnahme
    state.vosk_listening = False

    # Wenn kein persistenter Stream laeuft, eigenen oeffnen
    if not state.persistent_stream:
        try:
            state.stream = sd.InputStream(
                samplerate=SAMPLE_RATE, channels=1, dtype="float32",
                blocksize=int(SAMPLE_RATE * 0.1),
                device=state.audio_device, callback=persistent_audio_callback,
            )
            state.stream.start()
        except Exception as e:
            state.recording = False
            return

    asyncio.create_task(_auto_stop_timer())
    await broadcast({"type": "recording_started"})


async def create_session_internal(patient_name, location, medic, template_id):
    """Erstellt Session (intern, ohne HTTP)."""
    session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    state.sessions[session_id] = {
        "id": session_id,
        "template_id": template_id,
        "template_data": {},
        "patient_name": patient_name,
        "location": location,
        "medic": medic,
        "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "records": [],
        "created": datetime.now().isoformat(),
    }
    state.active_session = session_id
    await broadcast({"type": "session_created", "session": state.sessions[session_id]})


# ---------------------------------------------------------------------------
# API Routes
# ---------------------------------------------------------------------------
@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.get("/api/status")
async def get_status():
    return {
        "system": get_system_stats(),
        "model": state.current_model,
        "model_loaded": state.model_loaded,
        "model_loading": state.model_loading,
        "model_ram_mb": state.model_ram_mb,
        "recording": state.recording,
        "transcribing": state.transcribing,
        "audio_device": state.audio_device,
        "active_session": state.active_session,
        "session_data": state.sessions.get(state.active_session) if state.active_session else None,
        "language": state.language,
        "vosk_enabled": state.vosk_enabled,
        "vosk_listening": state.vosk_listening,
    }


@app.get("/api/templates")
async def get_templates():
    """Gibt alle verfuegbaren Templates zurueck."""
    return {"templates": list(RECORD_TEMPLATES.values())}


@app.get("/api/models")
async def get_models():
    return {"models": state.available_models(), "current": state.current_model}


@app.post("/api/models/load")
async def load_model(body: dict):
    name = body.get("model", "medium")
    path = MODELS_DIR / f"ggml-{name}.bin"
    if not path.exists():
        return {"error": f"Modell nicht gefunden: {path}"}

    if state.model_loading:
        return {"error": "Modell wird bereits geladen"}

    state.model_loading = True
    await broadcast({"type": "model_loading", "model": name})

    loop = asyncio.get_event_loop()
    success = await loop.run_in_executor(None, start_whisper_server, path)

    state.model_loading = False

    if success:
        state.model_path = path
        state.current_model = name
        await broadcast({
            "type": "model_loaded",
            "model": name,
            "ram_mb": state.model_ram_mb,
            "loaded": True,
        })
        return {"status": "ok", "model": name, "ram_mb": state.model_ram_mb}
    else:
        return {"error": f"Modell {name} konnte nicht geladen werden"}


@app.post("/api/models/unload")
async def unload_model():
    stop_whisper_server()
    old_model = state.current_model
    state.current_model = None
    state.model_path = None
    await broadcast({"type": "model_unloaded", "model": old_model})
    return {"status": "ok"}


@app.get("/api/devices")
async def get_devices():
    return {"devices": state.audio_devices(), "current": state.audio_device}


@app.post("/api/devices/select")
async def select_device(body: dict):
    state.audio_device = body.get("device_id")
    # Persistenten Stream mit neuem Device neu starten
    if state.persistent_stream:
        start_persistent_stream()
    return {"status": "ok", "device_id": state.audio_device}


@app.post("/api/language")
async def set_language(body: dict):
    state.language = body.get("language", "de")
    return {"status": "ok", "language": state.language}


@app.post("/api/session/create")
async def create_session(body: dict):
    template_id = body.get("template_id", "freitext")
    await create_session_internal(
        body.get("patient_name", "Unbekannt"),
        body.get("location", "Feld"),
        body.get("medic", ""),
        template_id,
    )
    return {"status": "ok", "session_id": state.active_session}


@app.post("/api/session/template-data")
async def save_template_data(body: dict):
    """Speichert Template-Felddaten fuer die aktive Session."""
    sid = body.get("session_id", state.active_session)
    if not sid or sid not in state.sessions:
        return {"error": "Keine aktive Session"}
    data = body.get("data", {})
    state.sessions[sid]["template_data"].update(data)
    return {"status": "ok"}


@app.get("/api/sessions")
async def list_sessions():
    return {"sessions": list(state.sessions.values()), "active": state.active_session}


@app.post("/api/session/select")
async def select_session(body: dict):
    sid = body.get("session_id")
    if sid in state.sessions:
        state.active_session = sid
        return {"status": "ok"}
    return {"error": "Session nicht gefunden"}


MAX_RECORD_SECONDS = 300
CHUNK_SECONDS = 25


@app.post("/api/record/start")
async def start_recording():
    if state.recording:
        return {"error": "Aufnahme laeuft bereits"}
    if not state.model_path:
        return {"error": "Kein Modell geladen"}

    state.audio_chunks = []
    state.recording = True
    state.vosk_listening = False

    # Wenn kein persistenter Stream, eigenen oeffnen
    if not state.persistent_stream:
        try:
            state.stream = sd.InputStream(
                samplerate=SAMPLE_RATE, channels=1, dtype="float32",
                blocksize=int(SAMPLE_RATE * 0.1),
                device=state.audio_device, callback=persistent_audio_callback,
            )
            state.stream.start()
        except Exception as e:
            state.recording = False
            return {"error": f"Audio-Fehler: {e}"}

    asyncio.create_task(_auto_stop_timer())
    await broadcast({"type": "recording_started"})
    return {"status": "recording", "max_seconds": MAX_RECORD_SECONDS}


async def _auto_stop_timer():
    await asyncio.sleep(MAX_RECORD_SECONDS)
    if state.recording:
        await broadcast({"type": "recording_auto_stop"})
        await stop_recording()


@app.post("/api/record/stop")
async def stop_recording():
    if not state.recording:
        return {"error": "Keine Aufnahme aktiv"}

    state.recording = False

    # Wenn eigener Stream (kein persistenter), schliessen
    if state.stream and not state.persistent_stream:
        state.stream.stop()
        state.stream.close()
        state.stream = None

    # Vosk wieder aktivieren
    if state.vosk_enabled:
        state.vosk_listening = True

    if not state.audio_chunks:
        return {"error": "Keine Audio-Daten"}

    audio = np.concatenate(state.audio_chunks, axis=0)
    total_duration = len(audio) / SAMPLE_RATE

    if total_duration < 0.5:
        await broadcast({"type": "recording_stopped", "duration": total_duration, "error": "Zu kurz"})
        return {"error": "Aufnahme zu kurz"}

    await broadcast({"type": "recording_stopped", "duration": round(total_duration, 1)})

    chunk_samples = CHUNK_SECONDS * SAMPLE_RATE
    chunks = []
    for i in range(0, len(audio), chunk_samples):
        chunk = audio[i:i + chunk_samples]
        if len(chunk) >= SAMPLE_RATE * 0.5:
            chunks.append(chunk)

    total_chunks = len(chunks)
    all_texts = []
    total_proc_time = 0

    state.transcribing = True
    await broadcast({"type": "transcribing", "chunks": total_chunks})

    loop = asyncio.get_event_loop()

    for idx, chunk in enumerate(chunks):
        await broadcast({
            "type": "transcribing_progress",
            "chunk": idx + 1,
            "total": total_chunks,
        })

        result = await loop.run_in_executor(None, run_transcribe, chunk, state.language)

        if result.get("error"):
            await broadcast({"type": "transcription_error", "error": result["error"], "chunk": idx + 1})
            continue

        text = result["text"]
        if text and text != "(Stille / nicht erkannt)":
            all_texts.append(text)
            total_proc_time += result.get("processing_time", 0)

            await broadcast({
                "type": "transcription_partial",
                "text": text,
                "chunk": idx + 1,
                "total": total_chunks,
            })

    state.transcribing = False

    # Vosk wieder aktiv
    if state.vosk_enabled:
        state.vosk_listening = True

    if not all_texts:
        await broadcast({"type": "transcription_error", "error": "Keine Sprache erkannt"})
        return {"error": "Keine Sprache erkannt"}

    full_text = " ".join(all_texts)
    rtf = total_proc_time / total_duration if total_duration > 0 else 0

    record_entry = {
        "time": datetime.now().strftime("%H:%M:%S"),
        "text": full_text,
        "audio_duration": round(total_duration, 2),
        "processing_time": round(total_proc_time, 2),
        "rtf": round(rtf, 3),
    }

    if state.active_session and state.active_session in state.sessions:
        state.sessions[state.active_session]["records"].append(record_entry)

    await broadcast({
        "type": "transcription_result",
        "record": record_entry,
        "session_id": state.active_session,
    })

    return {"status": "ok", "result": record_entry}


@app.post("/api/record/delete")
async def delete_record(body: dict):
    sid = body.get("session_id", state.active_session)
    idx = body.get("index")
    if sid in state.sessions and idx is not None:
        records = state.sessions[sid]["records"]
        if 0 <= idx < len(records):
            records.pop(idx)
            return {"status": "ok"}
    return {"error": "Eintrag nicht gefunden"}


@app.post("/api/session/analyze")
async def analyze_session(body: dict = None):
    """LLM analysiert Transkripte und extrahiert Template-Felder."""
    sid = body.get("session_id", state.active_session) if body else state.active_session
    if not sid or sid not in state.sessions:
        return {"error": "Keine aktive Session"}

    session = state.sessions[sid]
    template_id = session.get("template_id", "freitext")

    if template_id == "freitext":
        return {"status": "ok", "message": "Freitext braucht keine Analyse", "data": {}}

    if not session["records"]:
        return {"error": "Keine Transkripte vorhanden"}

    # Alle Transkripte zusammenfuegen
    full_text = " ".join(r["text"] for r in session["records"])

    await broadcast({"type": "analyzing", "session_id": sid})

    loop = asyncio.get_event_loop()
    extracted = await loop.run_in_executor(None, run_llm_extraction, template_id, full_text)

    if extracted:
        session["template_data"].update(extracted)
        await broadcast({
            "type": "analysis_complete",
            "session_id": sid,
            "data": extracted,
            "fields_count": len([v for v in extracted.values() if v]),
        })
        return {"status": "ok", "data": extracted}
    else:
        await broadcast({"type": "analysis_error", "error": "Extraktion fehlgeschlagen"})
        return {"error": "LLM-Extraktion fehlgeschlagen"}


@app.post("/api/export/docx")
async def export_docx(body: dict = None):
    sid = body.get("session_id", state.active_session) if body else state.active_session
    if not sid or sid not in state.sessions:
        return {"error": "Keine aktive Session"}

    session = state.sessions[sid]
    if not session["records"] and not session.get("template_data"):
        return {"error": "Keine Eintraege vorhanden"}

    filepath = generate_docx(session)
    return FileResponse(
        str(filepath),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filepath.name,
    )


@app.get("/api/files")
async def list_files():
    files = []
    for f in sorted(PROTOCOLS_DIR.glob("*")):
        files.append({
            "name": f.name,
            "size_kb": round(f.stat().st_size / 1024, 1),
            "modified": datetime.fromtimestamp(f.stat().st_mtime).isoformat(),
        })
    return {"files": files}


@app.get("/api/files/download/{filename}")
async def download_file(filename: str):
    filepath = PROTOCOLS_DIR / filename
    if not filepath.exists() or not filepath.is_relative_to(PROTOCOLS_DIR):
        return {"error": "Datei nicht gefunden"}
    return FileResponse(str(filepath), filename=filename)


@app.delete("/api/files/{filename}")
async def delete_file(filename: str):
    filepath = PROTOCOLS_DIR / filename
    if filepath.exists() and filepath.is_relative_to(PROTOCOLS_DIR):
        filepath.unlink()
        return {"status": "ok"}
    return {"error": "Datei nicht gefunden"}


# Vosk Toggle
@app.post("/api/vosk/toggle")
async def toggle_vosk(body: dict = None):
    if not state.vosk_model:
        return {"error": "Vosk nicht verfuegbar"}
    enable = body.get("enabled", not state.vosk_enabled) if body else not state.vosk_enabled
    state.vosk_enabled = enable
    if enable and not state.persistent_stream:
        start_persistent_stream()
    if enable:
        state.vosk_listening = not state.recording
    else:
        state.vosk_listening = False
        if state.persistent_stream and not state.recording:
            stop_persistent_stream()
    await broadcast({"type": "vosk_status", "enabled": state.vosk_enabled, "listening": state.vosk_listening})
    return {"status": "ok", "enabled": state.vosk_enabled}


@app.get("/api/vosk/status")
async def vosk_status():
    return {
        "enabled": state.vosk_enabled,
        "listening": state.vosk_listening,
        "available": state.vosk_model is not None,
    }


# ---------------------------------------------------------------------------
# WebSocket for real-time updates
# ---------------------------------------------------------------------------
@app.websocket("/ws")
async def websocket_endpoint(ws: WebSocket):
    await ws.accept()
    state.ws_clients.append(ws)

    await ws.send_json({
        "type": "init",
        "model": state.current_model,
        "model_loaded": state.model_loaded,
        "model_ram_mb": state.model_ram_mb,
        "recording": state.recording,
        "active_session": state.active_session,
        "session_data": state.sessions.get(state.active_session) if state.active_session else None,
        "vosk_enabled": state.vosk_enabled,
        "vosk_listening": state.vosk_listening,
    })

    try:
        while True:
            try:
                loop = asyncio.get_event_loop()
                stats = await loop.run_in_executor(None, get_system_stats)
                stats["model_loaded"] = state.model_loaded
                stats["model_ram_mb"] = state.model_ram_mb
                await ws.send_json({"type": "system_stats", "stats": stats})
            except Exception as e:
                print(f"WS stats error: {e}")

            try:
                if state.recording and state.audio_chunks:
                    last_chunk = state.audio_chunks[-1]
                    rms = float(np.sqrt(np.mean(last_chunk ** 2)))
                    duration = len(state.audio_chunks) * 0.1
                    await ws.send_json({
                        "type": "audio_level",
                        "rms": round(rms, 4),
                        "duration": round(duration, 1),
                    })
            except Exception:
                pass

            try:
                msg = await asyncio.wait_for(ws.receive_text(), timeout=2.0)
                data = json.loads(msg)
                if data.get("type") == "ping":
                    await ws.send_json({"type": "pong"})
            except asyncio.TimeoutError:
                pass
            except WebSocketDisconnect:
                break
            except Exception:
                break

    except WebSocketDisconnect:
        pass
    except Exception as e:
        print(f"WS error: {e}")
    finally:
        if ws in state.ws_clients:
            state.ws_clients.remove(ws)


# ---------------------------------------------------------------------------
# Startup / Shutdown
# ---------------------------------------------------------------------------
@app.on_event("startup")
async def startup():
    state.event_loop = asyncio.get_event_loop()

    # Vosk initialisieren
    init_vosk()

    # Whisper-Modell laden
    models = state.available_models()
    if models:
        best = next((m for m in models if m["name"] == "small"), models[0])
        path = MODELS_DIR / f"ggml-{best['name']}.bin"
        print(f"Lade Modell: {best['name']} ({best['size_mb']} MB)...")
        success = start_whisper_server(path)
        if success:
            state.model_path = path
            state.current_model = best["name"]
            print(f"Modell bereit: {best['name']} (~{state.model_ram_mb} MB RAM)")
        else:
            print("WARNUNG: Modell konnte nicht geladen werden!")

    # Persistenten Audio-Stream starten (fuer Vosk)
    if state.vosk_enabled:
        start_persistent_stream()

    # Vosk Command-Processor starten
    asyncio.create_task(process_vosk_commands())


@app.on_event("shutdown")
async def shutdown():
    stop_persistent_stream()
    stop_whisper_server()


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8080)
