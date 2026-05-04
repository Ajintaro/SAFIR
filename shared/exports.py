"""Export-Funktionen für die Patientendatenbank (DOCX/PDF/JSON/XML).

Generisch parametrisiert — nimmt `patients` als Liste, `device_id` und
`unit_name` als Strings, `output_dir` als Path. Wird sowohl vom Jetson-
Backend (`app.py`) als auch vom Surface-Backend (`backend/app.py`)
importiert, damit die Export-Logik an EINER Stelle lebt.

Die 4 Haupt-Funktionen:
    generate_json(patients, device_id, unit_name) -> bytes
    generate_xml(patients, device_id, unit_name) -> bytes
    generate_docx(patients, device_id, unit_name, output_dir) -> Path
    generate_pdf(patients, device_id, unit_name, output_dir) -> Path

JSON + XML geben Bytes zurück (direkt streamen), DOCX + PDF schreiben
in Dateien im output_dir und geben den Pfad zurück.
"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path


# ---------------------------------------------------------------------------
# Code-Mapping-Tabellen — analog zu den JS-Konstanten in templates/index.html.
# Werden fuer die Klartext-Aufschluesselung der 9-Liner-Codes in DOCX + PDF
# genutzt (z.B. Code "B" bei Line 3 -> "URGENT SURGICAL").
# ---------------------------------------------------------------------------
NINE_LINER_LABELS_EN = {
    "line1": "LINE 1 - LOCATION OF PICKUP SITE",
    "line2": "LINE 2 - RADIO FREQUENCY / CALL SIGN",
    "line3": "LINE 3 - NUMBER OF PATIENTS BY PRECEDENCE",
    "line4": "LINE 4 - SPECIAL EQUIPMENT REQUIRED",
    "line5": "LINE 5 - NUMBER OF PATIENTS BY TYPE",
    "line6": "LINE 6 - SECURITY OF PICKUP SITE",
    "line7": "LINE 7 - METHOD OF MARKING PICKUP SITE",
    "line8": "LINE 8 - PATIENT NATIONALITY AND STATUS",
    "line9": "LINE 9 - NBC / CBRN CONTAMINATION",
}
NINE_LINER_LABELS_DE = {
    "line1": "ZEILE 1 - AUFNAHMESTELLE",
    "line2": "ZEILE 2 - FUNKKANAL / RUFNAME",
    "line3": "ZEILE 3 - ANZAHL PATIENTEN NACH PRIORITAET",
    "line4": "ZEILE 4 - SONDERAUSRUESTUNG",
    "line5": "ZEILE 5 - ANZAHL PATIENTEN NACH TRANSPORTART",
    "line6": "ZEILE 6 - SICHERHEIT DER AUFNAHMESTELLE",
    "line7": "ZEILE 7 - MARKIERUNG DER AUFNAHMESTELLE",
    "line8": "ZEILE 8 - NATIONALITAET / STATUS",
    "line9": "ZEILE 9 - CBRN-KONTAMINATION",
}
PRECEDENCE_EN = {"A": "URGENT", "B": "URGENT SURGICAL", "C": "PRIORITY", "D": "ROUTINE", "E": "CONVENIENCE"}
PRECEDENCE_DE = {"A": "DRINGEND", "B": "DRINGEND CHIRURGISCH", "C": "PRIORITAET", "D": "ROUTINE", "E": "AUFSCHIEBBAR"}
SPECIAL_EQ_EN = {"A": "NONE", "B": "HOIST", "C": "EXTRACTION EQUIPMENT", "D": "VENTILATOR"}
SPECIAL_EQ_DE = {"A": "KEINE", "B": "WINDE", "C": "BERGEAUSRUESTUNG", "D": "BEATMUNGSGERAET"}
TYPE_EN = {"L": "LITTER", "A": "AMBULATORY"}
TYPE_DE = {"L": "TRAGEND", "A": "GEHFAEHIG"}
SECURITY_EN = {"N": "NO ENEMY TROOPS IN AREA", "P": "POSSIBLE ENEMY TROOPS IN AREA",
               "E": "ENEMY TROOPS IN AREA, APPROACH WITH CAUTION", "X": "ENEMY IN AREA - ARMED ESCORT REQUIRED"}
SECURITY_DE = {"N": "KEINE FEINDLICHEN KRAEFTE", "P": "MOEGLICHE FEINDLICHE KRAEFTE",
               "E": "FEINDLICHE KRAEFTE - VORSICHT", "X": "FEIND VOR ORT - BEWAFFNETE ESKORTE NOETIG"}
MARKING_EN = {"A": "PANEL", "B": "PYROTECHNIC SIGNAL", "C": "SMOKE SIGNAL", "D": "NONE", "E": "OTHER"}
MARKING_DE = {"A": "PANEL", "B": "PYROTECHNIK", "C": "RAUCH", "D": "KEINE", "E": "ANDERE"}
NATIONALITY_EN = {"A": "US MILITARY", "B": "US CIVILIAN", "C": "NON-US MILITARY",
                  "D": "NON-US CIVILIAN", "E": "EPW (ENEMY PRISONER OF WAR)"}
NATIONALITY_DE = {"A": "US-MILITAER", "B": "US-ZIVIL", "C": "NICHT-US-MILITAER (z.B. BUNDESWEHR)",
                  "D": "NICHT-US-ZIVIL", "E": "KRIEGSGEFANGENER"}
NBC_EN = {"N": "NUCLEAR", "B": "BIOLOGICAL", "C": "CHEMICAL", "R": "RADIOLOGICAL"}
NBC_DE = {"N": "NUKLEAR", "B": "BIOLOGISCH", "C": "CHEMISCH", "R": "RADIOLOGISCH"}


def _parse_count_codes(s: str) -> dict:
    """Parst '1 B' / 'B1' / 'B1 C2' etc. zu {'B': 1, 'C': 2}."""
    out: dict = {}
    if not s:
        return out
    tokens = s.replace(",", " ").replace(";", " ").upper().replace("BREAK", " ").split()
    i = 0
    while i < len(tokens):
        t = tokens[i]
        # Letter+Digit (B1)
        if len(t) >= 2 and t[0].isalpha() and t[1:].isdigit():
            out[t[0]] = out.get(t[0], 0) + int(t[1:])
        # Digit+Letter (1B)
        elif len(t) >= 2 and t[0].isdigit() and t[-1].isalpha():
            try:
                out[t[-1]] = out.get(t[-1], 0) + int(t[:-1])
            except ValueError:
                pass
        # Pure digit followed by letter
        elif t.isdigit() and i + 1 < len(tokens) and tokens[i + 1].isalpha() and len(tokens[i + 1]) == 1:
            out[tokens[i + 1]] = out.get(tokens[i + 1], 0) + int(t)
            i += 1
        # Single letter alone
        elif len(t) == 1 and t.isalpha():
            out[t] = out.get(t, 0) + 1
        i += 1
    return out


def expand_nine_liner_value(key: str, raw: str, lang: str = "en") -> list[str]:
    """Liefert eine Liste von Klartext-Zeilen fuer eine 9-Liner-Zeile.

    Beispiel: key='line3', raw='1 B', lang='en'
        -> ['A - URGENT: 0', 'B - URGENT SURGICAL: 1', 'C - PRIORITY: 0',
            'D - ROUTINE: 0', 'E - CONVENIENCE: 0']

    Bei Zeilen 1, 2 (Freitext) wird der Wert 1:1 als Single-Item-Liste
    zurueckgegeben.
    """
    is_en = lang == "en"
    val = (raw or "").strip()
    if not val:
        return ["—"]

    if key == "line3":
        m = PRECEDENCE_EN if is_en else PRECEDENCE_DE
        c = _parse_count_codes(val)
        return [f"{k} - {m[k]}: {c.get(k, 0)}" for k in ("A", "B", "C", "D", "E")]
    if key == "line4":
        m = SPECIAL_EQ_EN if is_en else SPECIAL_EQ_DE
        code = val[0].upper() if val else ""
        return [f"{code} - {m[code]}"] if code in m else [val]
    if key == "line5":
        m = TYPE_EN if is_en else TYPE_DE
        c = _parse_count_codes(val)
        return [f"{k} - {m[k]}: {c.get(k, 0)}" for k in ("L", "A")]
    if key == "line6":
        m = SECURITY_EN if is_en else SECURITY_DE
        code = val[0].upper() if val else ""
        return [f"{code} - {m[code]}"] if code in m else [val]
    if key == "line7":
        m = MARKING_EN if is_en else MARKING_DE
        code = val[0].upper() if val else ""
        return [f"{code} - {m[code]}"] if code in m else [val]
    if key == "line8":
        m = NATIONALITY_EN if is_en else NATIONALITY_DE
        code = val[0].upper() if val else ""
        # Multi-line value (z.B. "C - NON-US MILITARY\nDetails: 1 x German soldier")
        rest = val.split("\n", 1)
        extra = rest[1].strip() if len(rest) > 1 else ""
        if code in m:
            out = [f"{code} - {m[code]}"]
            if extra:
                out.append(extra)
            return out
        return [val]
    if key == "line9":
        m = NBC_EN if is_en else NBC_DE
        # Negation-Heuristik: "no known", "none", "keine"
        low = val.lower()
        if any(w in low for w in ("no known", "no contamination", "keine bekannte", "keine cbrn")):
            return ["No known nuclear, biological, chemical, or CBRN contamination."
                    if is_en else "Keine bekannte CBRN-Kontamination."]
        code = val[0].upper() if val else ""
        return [f"{code} - {m[code]}"] if code in m else [val]
    # line1, line2: Freitext
    return [val]


def detect_nine_liner_lang(patient: dict) -> str:
    """Heuristik: 'en' wenn p.language='en' oder line1/line2 englische Marker
    enthalten, sonst 'de'."""
    lang = (patient or {}).get("language", "")
    if lang == "en":
        return "en"
    if lang == "de":
        return "de"
    nl = patient.get("nine_liner") or {}
    blob = (nl.get("line1", "") + " " + nl.get("line2", "")).lower()
    if any(m in blob for m in ("medical net", "call sign", "pickup", "casualty")):
        return "en"
    return "de"


# FMC-Sektionen-Labels (englisch — FMC ist NATO-Standard, immer EN)
FMC_SECTION_LABELS = {
    "section_a": "A - IDENTIFICATION",
    "section_b": "B - CAUSE",
    "section_c": "C - INITIAL ASSESSMENT",
    "section_d": "D - VITAL SIGNS",
    "section_e": "E - TREATMENT",
    "section_f": "F - MOVEMENT / EVACUATION",
    "section_g": "G - DOCUMENTATION",
}
FMC_FIELD_LABELS = {
    # Section A
    "last_name": "Last Name", "first_name": "First Name", "rank": "Rank",
    "sex": "Sex", "dob": "Date of Birth", "service_number": "Service Number",
    "nationality": "Nationality / Armed Forces", "unit": "Unit of Origin",
    # Section B
    "casualty_type": "Casualty Type", "datetime_injury": "Date / Time of Injury",
    "mechanism": "Mechanism of Injury",
    # Section C
    "time_first_assessment": "Time of First Assessment",
    "general_condition": "General Condition", "airway": "Airway",
    "breathing": "Breathing", "chest": "Chest",
    "main_injury": "Main Injury", "main_injury_region": "Injury Region",
    "additional_injury": "Additional Injury",
    "loss_of_consciousness": "Loss of Consciousness",
    "allergies": "Allergies", "blood_group": "Blood Group",
    # Section E
    "tourniquet": "Tourniquet", "tourniquet_time": "Tourniquet Time",
    "hemorrhage_control": "Hemorrhage Control",
    "hypothermia_prevention": "Hypothermia Prevention",
    "iv_access": "IV Access",
    "immobilization": "Immobilization",
    "surgical_procedure": "Surgical Procedure",
    # Section F
    "evacuation_priority": "Evacuation Priority",
    "transport_category": "Transport Category",
    "destination": "Recommended Destination",
    # Section G
    "recorded_by": "Recorded By", "function": "Function",
}

ATMIST_LABELS = {
    "line1": "A - Angaben / Alter",
    "line2": "T - Time (Zeitpunkt)",
    "line3": "M - Mechanismus",
    "line4": "I - Injury (Verletzungen)",
    "line5": "S - Signs (Vitals)",
    "line6": "T - Treatment (Massnahmen)",
}


# ---------------------------------------------------------------------------
# JSON
# ---------------------------------------------------------------------------
def generate_json(patients: list, device_id: str, unit_name: str) -> bytes:
    """Komplette Patientendatenbank als JSON-Bytes. Schema-versioniert."""
    payload = {
        "schema_version": "1.0",
        "exported_at": datetime.now().isoformat(),
        "device_id": device_id,
        "unit_name": unit_name,
        "patient_count": len(patients),
        "patients": list(patients),
    }
    return json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")


# ---------------------------------------------------------------------------
# XML
# ---------------------------------------------------------------------------
def _xml_escape(val) -> str:
    """Minimaler XML-Text-Escape — & < > nur. Werte landen in Element-
    Text, nicht in Attribut-Werten."""
    if val is None:
        return ""
    s = str(val)
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _build_patient_xml(patient: dict) -> str:
    """Konvertiert einen Patient-Dict in einen XML-Block. Rekursiv für
    Dicts (verschachtelt) und Listen (`<item>`-Wrapper)."""
    def _emit(key: str, value, indent: int) -> str:
        pad = "  " * indent
        if isinstance(value, dict):
            if not value:
                return f'{pad}<{key}/>\n'
            inner = "".join(_emit(k, v, indent + 1) for k, v in value.items())
            return f'{pad}<{key}>\n{inner}{pad}</{key}>\n'
        if isinstance(value, (list, tuple)):
            if not value:
                return f'{pad}<{key}/>\n'
            inner = "".join(_emit("item", v, indent + 1) for v in value)
            return f'{pad}<{key}>\n{inner}{pad}</{key}>\n'
        return f'{pad}<{key}>{_xml_escape(value)}</{key}>\n'

    pid = patient.get("patient_id", "unknown")
    status = patient.get("flow_status", "")
    out = [f'  <patient id="{_xml_escape(pid)}" status="{_xml_escape(status)}">\n']
    for k, v in patient.items():
        if k in ("patient_id", "flow_status"):
            continue  # schon als Attribut
        out.append(_emit(k, v, 2))
    out.append("  </patient>\n")
    return "".join(out)


def generate_xml(patients: list, device_id: str, unit_name: str) -> bytes:
    """Komplette Patientendatenbank als XML-Bytes. Format:
    <safir-export schema="1.0" device-id="..." unit-name="..." ...>
      <patients>
        <patient id="..." status="..."><name>...</name>...</patient>
      </patients>
    </safir-export>"""
    now_iso = datetime.now().isoformat()
    parts = ['<?xml version="1.0" encoding="UTF-8"?>\n']
    parts.append(
        f'<safir-export schema="1.0" device-id="{_xml_escape(device_id)}" '
        f'unit-name="{_xml_escape(unit_name)}" exported="{now_iso}" '
        f'patient-count="{len(patients)}">\n'
    )
    parts.append("  <patients>\n")
    for patient in patients:
        parts.append(_build_patient_xml(patient))
    parts.append("  </patients>\n")
    parts.append("</safir-export>\n")
    return "".join(parts).encode("utf-8")


# ---------------------------------------------------------------------------
# DOCX (via python-docx)
# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# DOCX-Helper fuer FMC-Sektionen (werden von generate_docx aufgerufen)
# ---------------------------------------------------------------------------
def _render_fmc_section_docx(doc, section_key: str, data: dict) -> None:
    """Rendert eine flache FMC-Sektion (A, B, C, F, G) als 2-Spalten-Tabelle."""
    if not data:
        return
    doc.add_heading(FMC_SECTION_LABELS[section_key], level=3)
    # Felder in der gewuenschten Reihenfolge — andere Felder ans Ende
    preferred_order = [
        # Section A
        "last_name", "first_name", "rank", "sex", "dob", "service_number",
        "nationality", "unit",
        # Section B
        "casualty_type", "datetime_injury", "mechanism",
        # Section C
        "time_first_assessment", "general_condition", "airway", "breathing",
        "chest", "main_injury", "main_injury_region", "additional_injury",
        "loss_of_consciousness", "allergies", "blood_group",
        # Section F
        "evacuation_priority", "transport_category", "destination",
        # Section G
        "recorded_by", "function",
    ]
    keys = [k for k in preferred_order if k in data]
    keys += [k for k in data.keys() if k not in keys]
    if not keys:
        return
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k in keys:
        v = data.get(k)
        # Listen wie 'considerations' als Bullet-String
        if isinstance(v, list):
            v_str = "\n".join(f"- {item}" for item in v) if v else "—"
        elif v in (None, ""):
            v_str = "—"
        else:
            v_str = str(v)
        row = t.add_row().cells
        row[0].text = FMC_FIELD_LABELS.get(k, k.replace("_", " ").title())
        for pr in row[0].paragraphs:
            for r in pr.runs:
                r.bold = True
        row[1].text = v_str
    # Considerations-Liste (nur Section F)
    if section_key == "section_f" and isinstance(data.get("considerations"), list) and data["considerations"]:
        # ist oben schon mit drin, Skip
        pass
    doc.add_paragraph()


def _render_fmc_vitals_docx(doc, vitals_list: list) -> None:
    """Section D: Vital Signs als Tabelle mit allen Messpunkten."""
    if not vitals_list:
        return
    doc.add_heading(FMC_SECTION_LABELS["section_d"], level=3)
    headers = ["Time", "Pulse", "BP", "Resp", "SpO2", "GCS", "Pain"]
    keys = ["time", "pulse", "bp", "resp_rate", "spo2", "gcs", "pain"]
    t = doc.add_table(rows=1 + len(vitals_list), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for pr in c.paragraphs:
            for r in pr.runs:
                r.bold = True
    for ri, v in enumerate(vitals_list, start=1):
        if not isinstance(v, dict):
            continue
        for ci, k in enumerate(keys):
            t.rows[ri].cells[ci].text = str(v.get(k, "") or "—")
    doc.add_paragraph()


def _render_fmc_treatment_docx(doc, e: dict) -> None:
    """Section E: Treatment — Tourniquet, Hemorrhage Control, Meds, Fluids."""
    if not e:
        return
    doc.add_heading(FMC_SECTION_LABELS["section_e"], level=3)
    # Erst Schluessel-Felder (Tourniquet, Hemorrhage etc.) als KV-Tabelle
    kv = []
    for key in ("tourniquet", "tourniquet_time", "hemorrhage_control",
                "hypothermia_prevention", "iv_access", "immobilization",
                "surgical_procedure"):
        v = e.get(key)
        if v:
            kv.append((FMC_FIELD_LABELS.get(key, key), str(v)))
    if kv:
        t = doc.add_table(rows=len(kv), cols=2)
        t.style = "Table Grid"
        for i, (label, val) in enumerate(kv):
            c = t.rows[i].cells[0]
            c.text = label
            for pr in c.paragraphs:
                for r in pr.runs:
                    r.bold = True
            t.rows[i].cells[1].text = val
        doc.add_paragraph()
    # Medikamente
    meds = e.get("medications") or []
    if meds:
        doc.add_heading("Medications", level=4)
        mt = doc.add_table(rows=1 + len(meds), cols=4)
        mt.style = "Table Grid"
        for i, h in enumerate(["Name", "Dose", "Route", "Time"]):
            c = mt.rows[0].cells[i]
            c.text = h
            for pr in c.paragraphs:
                for r in pr.runs:
                    r.bold = True
        for ri, m in enumerate(meds, start=1):
            if not isinstance(m, dict):
                continue
            mt.rows[ri].cells[0].text = str(m.get("name") or m.get("drug") or "—")
            mt.rows[ri].cells[1].text = str(m.get("dose") or "—")
            mt.rows[ri].cells[2].text = str(m.get("route") or "—")
            mt.rows[ri].cells[3].text = str(m.get("time") or "—")
        doc.add_paragraph()
    # Fluids
    fluids = e.get("fluids") or []
    if fluids:
        doc.add_heading("Fluids", level=4)
        ft = doc.add_table(rows=1 + len(fluids), cols=4)
        ft.style = "Table Grid"
        for i, h in enumerate(["Name", "Volume", "Route", "Time"]):
            c = ft.rows[0].cells[i]
            c.text = h
            for pr in c.paragraphs:
                for r in pr.runs:
                    r.bold = True
        for ri, fl in enumerate(fluids, start=1):
            if not isinstance(fl, dict):
                continue
            ft.rows[ri].cells[0].text = str(fl.get("name") or "—")
            ft.rows[ri].cells[1].text = str(fl.get("volume") or "—")
            ft.rows[ri].cells[2].text = str(fl.get("route") or "—")
            ft.rows[ri].cells[3].text = str(fl.get("time") or "—")
        doc.add_paragraph()


def generate_docx(patients: list, device_id: str, unit_name: str,
                  output_dir: Path) -> Path:
    """Eine DOCX-Datei mit Übersichtstabelle + Detail-Block pro Patient.
    Schreibt nach output_dir/safir-patients-<timestamp>.docx.

    Raises ImportError wenn python-docx nicht installiert ist."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("SAFIR — Patientendatenbank Export")
    run.bold = True
    run.font.size = Pt(16)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(
        f"Exportiert am {datetime.now().strftime('%d.%m.%Y %H:%M')}  ·  "
        f"Geraet {device_id or '?'}  ·  "
        f"Einheit {unit_name or '?'}  ·  "
        f"{len(patients)} Patient(en)"
    )
    meta_run.font.size = Pt(9)
    meta_run.italic = True
    doc.add_paragraph()

    if not patients:
        doc.add_paragraph("(Keine Patienten in der Datenbank)")
    else:
        # Uebersichtstabelle
        doc.add_heading("Uebersicht", level=2)
        headers = ["Patient-ID", "Name", "Dienstgrad", "Triage", "Status", "Sync"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr_row = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_row[i].text = h
            for p in hdr_row[i].paragraphs:
                for r in p.runs:
                    r.bold = True
        for p in patients:
            row = table.add_row().cells
            row[0].text = p.get("patient_id", "")
            row[1].text = p.get("name", "Unbekannt") or "Unbekannt"
            row[2].text = p.get("rank", "") or "—"
            row[3].text = p.get("triage", "") or "—"
            row[4].text = p.get("flow_status", "") or p.get("status", "")
            row[5].text = "✓" if p.get("synced") else "—"
        doc.add_paragraph()

        # Detail-Sektion pro Patient
        for idx, p in enumerate(patients):
            doc.add_page_break()
            heading = doc.add_heading(
                f"Patient {idx + 1} — {p.get('name', 'Unbekannt')}", level=1
            )
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Stammdaten
            stamm = [
                ("Patient-ID", p.get("patient_id", "")),
                ("Dienstgrad", p.get("rank", "") or "—"),
                ("Einheit", p.get("unit", "") or "—"),
                ("Nationalitaet", p.get("nationality", "") or "—"),
                ("Blutgruppe", p.get("blood_type", "") or "—"),
                ("Allergien", p.get("allergies", "") or "—"),
                ("Triage", p.get("triage", "") or "—"),
                ("Status", p.get("status", "") or "—"),
                ("Flow-Status", p.get("flow_status", "") or "—"),
                ("Aktuelle Rolle", p.get("current_role", "") or "—"),
                ("RFID-UID", p.get("rfid_tag_id", "") or "—"),
                ("Erfasst von", p.get("created_by", "") or "—"),
                ("Erfasst am", p.get("timestamp_created", "") or "—"),
                ("An Leitstelle gemeldet", "Ja" if p.get("synced") else "Nein"),
            ]
            t = doc.add_table(rows=len(stamm), cols=2)
            t.style = "Table Grid"
            for i, (label, value) in enumerate(stamm):
                c0 = t.rows[i].cells[0]
                c0.text = label
                for pr in c0.paragraphs:
                    for r in pr.runs:
                        r.bold = True
                t.rows[i].cells[1].text = str(value)
            doc.add_paragraph()

            # 9-Liner MEDEVAC (NATO ATP-3.7.2 Vollformat, Anhang A)
            nl = p.get("nine_liner") or {}
            filled = [k for k in [f"line{i}" for i in range(1, 10)] if nl.get(k)]
            if p.get("template_type") == "9liner" or filled:
                lang = detect_nine_liner_lang(p)
                title = ("9-Line MEDEVAC Request" if lang == "en"
                         else "9-Liner MEDEVAC Anforderung")
                doc.add_heading(f"{title}  ({len(filled)}/9, NATO ATP-3.7.2)", level=2)
                labels = NINE_LINER_LABELS_EN if lang == "en" else NINE_LINER_LABELS_DE
                nt = doc.add_table(rows=len(labels), cols=2)
                nt.style = "Table Grid"
                for i, key in enumerate(labels.keys()):
                    c0 = nt.rows[i].cells[0]
                    c0.text = labels[key]
                    for pr in c0.paragraphs:
                        for r in pr.runs:
                            r.bold = True
                    expanded = expand_nine_liner_value(key, nl.get(key, ""), lang)
                    nt.rows[i].cells[1].text = "\n".join(expanded)
                if nl.get("remarks"):
                    nt.add_row()
                    rrow = nt.rows[-1].cells
                    rrow[0].text = "REMARKS"
                    for pr in rrow[0].paragraphs:
                        for r in pr.runs:
                            r.bold = True
                    rrow[1].text = nl["remarks"]
                doc.add_paragraph()

            # ATMIST Patientenuebergabe
            atm = p.get("atmist") or {}
            atm_filled = [k for k in [f"line{i}" for i in range(1, 7)] if atm.get(k)]
            if p.get("template_type") == "atmist" or atm_filled:
                doc.add_heading(f"ATMIST Patientenuebergabe  ({len(atm_filled)}/6)", level=2)
                at = doc.add_table(rows=len(ATMIST_LABELS), cols=2)
                at.style = "Table Grid"
                for i, key in enumerate(ATMIST_LABELS.keys()):
                    c0 = at.rows[i].cells[0]
                    c0.text = ATMIST_LABELS[key]
                    for pr in c0.paragraphs:
                        for r in pr.runs:
                            r.bold = True
                    at.rows[i].cells[1].text = str(atm.get(key, "") or "—")
                doc.add_paragraph()

            # NATO Field Medical Card (AMedP-8.1 Annex A)
            fmc = p.get("fmc") or {}
            if p.get("template_type") == "fmc" or fmc:
                doc.add_heading("NATO Field Medical Card  (AMedP-8.1 Annex A)", level=2)
                _render_fmc_section_docx(doc, "section_a", fmc.get("section_a") or {})
                _render_fmc_section_docx(doc, "section_b", fmc.get("section_b") or {})
                _render_fmc_section_docx(doc, "section_c", fmc.get("section_c") or {})
                _render_fmc_vitals_docx(doc, fmc.get("section_d") or [])
                _render_fmc_treatment_docx(doc, fmc.get("section_e") or {})
                _render_fmc_section_docx(doc, "section_f", fmc.get("section_f") or {})
                _render_fmc_section_docx(doc, "section_g", fmc.get("section_g") or {})
                doc.add_paragraph()

            # Vitals
            vitals = p.get("vitals") or {}
            if any(vitals.values()):
                doc.add_heading("Vitalwerte", level=2)
                vt = doc.add_table(rows=0, cols=2)
                vt.style = "Table Grid"
                vital_labels = {
                    "pulse": "Puls (bpm)",
                    "bp": "Blutdruck",
                    "resp_rate": "Atemfrequenz",
                    "spo2": "SpO2 (%)",
                    "temp": "Temperatur (°C)",
                    "gcs": "GCS",
                }
                for key, label in vital_labels.items():
                    val = vitals.get(key)
                    if val:
                        row = vt.add_row().cells
                        row[0].text = label
                        for pr in row[0].paragraphs:
                            for r in pr.runs:
                                r.bold = True
                        row[1].text = str(val)
                doc.add_paragraph()

            # Verletzungen
            injuries = p.get("injuries") or []
            if injuries:
                doc.add_heading("Verletzungen", level=2)
                for inj in injuries:
                    doc.add_paragraph(f"• {inj}")
                doc.add_paragraph()

            # Behandlungen / Medikamente
            treatments = p.get("treatments") or []
            medications = p.get("medications") or []
            if treatments or medications:
                doc.add_heading("Behandlungen / Medikamente", level=2)
                for item in treatments:
                    val = item if isinstance(item, str) else (
                        item.get("description") or json.dumps(item, ensure_ascii=False)
                    )
                    doc.add_paragraph(f"• {val}")
                for item in medications:
                    val = item if isinstance(item, str) else (
                        item.get("name") or json.dumps(item, ensure_ascii=False)
                    )
                    doc.add_paragraph(f"• {val}")
                doc.add_paragraph()

            # Transkripte
            transcripts = p.get("transcripts") or []
            if transcripts:
                doc.add_heading("Transkripte", level=2)
                for tr in transcripts:
                    if isinstance(tr, dict):
                        time_str = tr.get("time", "")
                        text = tr.get("text", "")
                    else:
                        time_str = ""
                        text = str(tr)
                    para = doc.add_paragraph()
                    if time_str:
                        rt = para.add_run(f"[{time_str}] ")
                        rt.bold = True
                        rt.font.size = Pt(9)
                    para.add_run(text)
                doc.add_paragraph()

            # Timeline (letzte 10)
            timeline = p.get("timeline") or []
            if timeline:
                doc.add_heading("Timeline", level=2)
                for ev in timeline[-10:]:
                    if isinstance(ev, dict):
                        line = (
                            f"{ev.get('time', '')}  ·  [{ev.get('role', '')}]  "
                            f"{ev.get('event', '')}  —  {ev.get('details', '')}"
                        )
                    else:
                        line = str(ev)
                    para = doc.add_paragraph(line)
                    for r in para.runs:
                        r.font.size = Pt(9)

    # Footer
    doc.add_paragraph()
    doc.add_heading("Unterschrift", level=2)
    doc.add_paragraph("_" * 30)

    output_dir.mkdir(exist_ok=True, parents=True)
    filename = f"safir-patients-{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = output_dir / filename
    doc.save(str(filepath))
    return filepath


# ---------------------------------------------------------------------------
# PDF-Helper fuer FMC-Sektionen (von generate_pdf aufgerufen)
# ---------------------------------------------------------------------------
def _render_fmc_section_pdf(story, styles, esc_fn, kv_table_fn,
                            section_key: str, data: dict) -> None:
    """Flache FMC-Sektion (A, B, C, F, G) als 2-Spalten-Tabelle in reportlab."""
    if not data:
        return
    from reportlab.platypus import Paragraph, Spacer
    story.append(Paragraph(FMC_SECTION_LABELS[section_key], styles["SAFIRH3"]))
    preferred_order = [
        "last_name", "first_name", "rank", "sex", "dob", "service_number",
        "nationality", "unit",
        "casualty_type", "datetime_injury", "mechanism",
        "time_first_assessment", "general_condition", "airway", "breathing",
        "chest", "main_injury", "main_injury_region", "additional_injury",
        "loss_of_consciousness", "allergies", "blood_group",
        "evacuation_priority", "transport_category", "destination",
        "considerations",
        "recorded_by", "function",
    ]
    keys = [k for k in preferred_order if k in data]
    keys += [k for k in data.keys() if k not in keys]
    rows = []
    for k in keys:
        v = data.get(k)
        if isinstance(v, list):
            v_str = "<br/>".join(f"- {esc_fn(item)}" for item in v) if v else "—"
        elif v in (None, ""):
            v_str = "—"
        else:
            v_str = esc_fn(str(v))
        label = FMC_FIELD_LABELS.get(k, k.replace("_", " ").title())
        rows.append([label, Paragraph(v_str, styles["Normal"])])
    if rows:
        story.append(kv_table_fn(rows))
        story.append(Spacer(1, 4))


def _render_fmc_vitals_pdf(story, styles, esc_fn, vitals_list: list) -> None:
    """Section D: Vital Signs als Mehrspalten-Tabelle in reportlab."""
    if not vitals_list:
        return
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors as _rl_colors
    from reportlab.lib.units import mm
    story.append(Paragraph(FMC_SECTION_LABELS["section_d"], styles["SAFIRH3"]))
    headers = ["Time", "Pulse", "BP", "Resp", "SpO2", "GCS", "Pain"]
    keys = ["time", "pulse", "bp", "resp_rate", "spo2", "gcs", "pain"]
    rows = [headers]
    for v in vitals_list:
        if not isinstance(v, dict):
            continue
        rows.append([str(v.get(k, "") or "—") for k in keys])
    t = Table(rows, colWidths=[20 * mm] + [16 * mm] * (len(headers) - 1))
    t.setStyle(TableStyle([
        ("FONT",       (0, 0), (-1, 0), "Helvetica-Bold", 9),
        ("FONT",       (0, 1), (-1, -1), "Helvetica", 9),
        ("BACKGROUND", (0, 0), (-1, 0), _rl_colors.HexColor("#e0d8b8")),
        ("GRID",       (0, 0), (-1, -1), 0.5, _rl_colors.black),
        ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(t)
    story.append(Spacer(1, 6))


def _render_fmc_treatment_pdf(story, styles, esc_fn, kv_table_fn, e: dict) -> None:
    """Section E: Treatment + Medications + Fluids."""
    if not e:
        return
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors as _rl_colors
    from reportlab.lib.units import mm
    story.append(Paragraph(FMC_SECTION_LABELS["section_e"], styles["SAFIRH3"]))
    kv_rows = []
    for key in ("tourniquet", "tourniquet_time", "hemorrhage_control",
                "hypothermia_prevention", "iv_access", "immobilization",
                "surgical_procedure"):
        v = e.get(key)
        if v:
            kv_rows.append([FMC_FIELD_LABELS.get(key, key),
                            Paragraph(esc_fn(str(v)), styles["Normal"])])
    if kv_rows:
        story.append(kv_table_fn(kv_rows))
        story.append(Spacer(1, 4))

    def _med_fluid_table(items, headers, keys):
        if not items:
            return None
        rows = [headers]
        for it in items:
            if not isinstance(it, dict):
                continue
            rows.append([str(it.get(k) or "—") for k in keys])
        if len(rows) <= 1:
            return None
        t = Table(rows, colWidths=[40 * mm, 25 * mm, 20 * mm, 20 * mm])
        t.setStyle(TableStyle([
            ("FONT",       (0, 0), (-1, 0), "Helvetica-Bold", 9),
            ("FONT",       (0, 1), (-1, -1), "Helvetica", 9),
            ("BACKGROUND", (0, 0), (-1, 0), _rl_colors.HexColor("#e0d8b8")),
            ("GRID",       (0, 0), (-1, -1), 0.5, _rl_colors.black),
        ]))
        return t

    meds = e.get("medications") or []
    if meds:
        story.append(Paragraph("Medications", styles["SAFIRH3"]))
        # Gemma liefert manchmal {drug: ..., dose: ...} statt {name: ..., dose: ...}
        norm_meds = [{**m, "name": (m.get("name") or m.get("drug") or "")} for m in meds if isinstance(m, dict)]
        t = _med_fluid_table(norm_meds, ["Name", "Dose", "Route", "Time"],
                             ["name", "dose", "route", "time"])
        if t:
            story.append(t)
            story.append(Spacer(1, 4))
    fluids = e.get("fluids") or []
    if fluids:
        story.append(Paragraph("Fluids", styles["SAFIRH3"]))
        t = _med_fluid_table(fluids, ["Name", "Volume", "Route", "Time"],
                             ["name", "volume", "route", "time"])
        if t:
            story.append(t)
            story.append(Spacer(1, 4))


# ---------------------------------------------------------------------------
# PDF (via reportlab)
# ---------------------------------------------------------------------------
def generate_pdf(patients: list, device_id: str, unit_name: str,
                 output_dir: Path) -> Path:
    """Eine PDF-Datei mit Übersicht + Detail pro Patient. Design in
    Bundeswehr-Olive für Brand-Konsistenz. Schreibt nach
    output_dir/safir-patients-<timestamp>.pdf.

    Raises ImportError wenn reportlab nicht installiert ist."""
    from reportlab.lib import colors as _rl_colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak,
    )

    # Bundeswehr-Olive-Akzent
    BW_TAN = _rl_colors.HexColor("#c8b878")
    BW_DARK = _rl_colors.HexColor("#3a4a2e")
    BW_BG = _rl_colors.HexColor("#f4f1e3")

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="SAFIRTitle", parent=styles["Title"],
        fontSize=18, textColor=BW_DARK, spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="SAFIRMeta", parent=styles["Normal"],
        fontSize=9, textColor=_rl_colors.grey, spaceAfter=16, alignment=1,
    ))
    styles.add(ParagraphStyle(
        name="SAFIRH1", parent=styles["Heading1"],
        fontSize=14, textColor=BW_DARK, spaceBefore=8, spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="SAFIRH2", parent=styles["Heading2"],
        fontSize=11, textColor=BW_DARK, spaceBefore=8, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="SAFIRH3", parent=styles["Heading3"],
        fontSize=10, textColor=BW_DARK, spaceBefore=6, spaceAfter=3,
    ))

    def _kv_table(rows: list, col_widths=(55 * mm, 105 * mm)):
        tbl = Table(rows, colWidths=col_widths, hAlign="LEFT")
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), BW_BG),
            ("TEXTCOLOR", (0, 0), (0, -1), BW_DARK),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("LINEBELOW", (0, 0), (-1, -1), 0.3, _rl_colors.lightgrey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return tbl

    def _esc(v):
        if v is None:
            return ""
        s = str(v)
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story = []
    story.append(Paragraph("SAFIR — Patientendatenbank Export", styles["SAFIRTitle"]))
    meta = (
        f"Exportiert am {datetime.now().strftime('%d.%m.%Y %H:%M')}  ·  "
        f"Gerät {_esc(device_id or '?')}  ·  "
        f"Einheit {_esc(unit_name or '?')}  ·  "
        f"{len(patients)} Patient(en)"
    )
    story.append(Paragraph(meta, styles["SAFIRMeta"]))

    if not patients:
        story.append(Paragraph("(Keine Patienten in der Datenbank)", styles["Normal"]))
    else:
        # Uebersichtstabelle
        story.append(Paragraph("Übersicht", styles["SAFIRH1"]))
        overview_rows = [["Patient-ID", "Name", "Dienstgrad", "Triage", "Status", "Sync"]]
        for p in patients:
            overview_rows.append([
                _esc(p.get("patient_id", "")),
                _esc(p.get("name", "Unbekannt") or "Unbekannt"),
                _esc(p.get("rank", "") or "—"),
                _esc(p.get("triage", "") or "—"),
                _esc(p.get("flow_status", "") or p.get("status", "")),
                "✓" if p.get("synced") else "—",
            ])
        ov_tbl = Table(overview_rows, hAlign="LEFT",
                       colWidths=(28 * mm, 40 * mm, 30 * mm, 18 * mm, 28 * mm, 14 * mm))
        ov_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BW_DARK),
            ("TEXTCOLOR", (0, 0), (-1, 0), _rl_colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [_rl_colors.white, BW_BG]),
            ("GRID", (0, 0), (-1, -1), 0.3, _rl_colors.lightgrey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(ov_tbl)

        # Detail-Sektion pro Patient
        for idx, p in enumerate(patients):
            story.append(PageBreak())
            story.append(Paragraph(
                f"Patient {idx + 1} — {_esc(p.get('name', 'Unbekannt'))}",
                styles["SAFIRH1"],
            ))

            stamm_rows = [
                ["Patient-ID", _esc(p.get("patient_id", ""))],
                ["Dienstgrad", _esc(p.get("rank", "") or "—")],
                ["Einheit", _esc(p.get("unit", "") or "—")],
                ["Nationalität", _esc(p.get("nationality", "") or "—")],
                ["Blutgruppe", _esc(p.get("blood_type", "") or "—")],
                ["Allergien", _esc(p.get("allergies", "") or "—")],
                ["Triage", _esc(p.get("triage", "") or "—")],
                ["Status", _esc(p.get("status", "") or "—")],
                ["Flow-Status", _esc(p.get("flow_status", "") or "—")],
                ["Aktuelle Rolle", _esc(p.get("current_role", "") or "—")],
                ["RFID-UID", _esc(p.get("rfid_tag_id", "") or "—")],
                ["Erfasst von", _esc(p.get("created_by", "") or "—")],
                ["Erfasst am", _esc(p.get("timestamp_created", "") or "—")],
                ["Gemeldet", "Ja" if p.get("synced") else "Nein"],
            ]
            story.append(_kv_table(stamm_rows))
            story.append(Spacer(1, 6))

            nl = p.get("nine_liner") or {}
            filled = [k for k in [f"line{i}" for i in range(1, 10)] if nl.get(k)]
            if p.get("template_type") == "9liner" or filled:
                lang = detect_nine_liner_lang(p)
                title = ("9-Line MEDEVAC Request" if lang == "en"
                         else "9-Liner MEDEVAC Anforderung")
                story.append(Paragraph(
                    f"{title} ({len(filled)}/9 - NATO ATP-3.7.2)",
                    styles["SAFIRH2"],
                ))
                labels = NINE_LINER_LABELS_EN if lang == "en" else NINE_LINER_LABELS_DE
                nl_rows = []
                for key in labels.keys():
                    expanded = expand_nine_liner_value(key, nl.get(key, ""), lang)
                    val_html = "<br/>".join(_esc(line) for line in expanded)
                    nl_rows.append([labels[key], Paragraph(val_html, styles["Normal"])])
                if nl.get("remarks"):
                    nl_rows.append(["REMARKS", Paragraph(_esc(nl["remarks"]), styles["Normal"])])
                story.append(_kv_table(nl_rows))
                story.append(Spacer(1, 6))

            # ATMIST Patientenuebergabe
            atm = p.get("atmist") or {}
            atm_filled = [k for k in [f"line{i}" for i in range(1, 7)] if atm.get(k)]
            if p.get("template_type") == "atmist" or atm_filled:
                story.append(Paragraph(
                    f"ATMIST Patientenuebergabe ({len(atm_filled)}/6)",
                    styles["SAFIRH2"],
                ))
                atm_rows = []
                for key in ATMIST_LABELS.keys():
                    val = atm.get(key, "") or "—"
                    atm_rows.append([ATMIST_LABELS[key],
                                     Paragraph(_esc(val), styles["Normal"])])
                story.append(_kv_table(atm_rows))
                story.append(Spacer(1, 6))

            # NATO Field Medical Card (AMedP-8.1 Annex A)
            fmc = p.get("fmc") or {}
            if p.get("template_type") == "fmc" or fmc:
                story.append(Paragraph(
                    "NATO Field Medical Card (AMedP-8.1 Annex A)",
                    styles["SAFIRH2"],
                ))
                _render_fmc_section_pdf(story, styles, _esc, _kv_table,
                                        "section_a", fmc.get("section_a") or {})
                _render_fmc_section_pdf(story, styles, _esc, _kv_table,
                                        "section_b", fmc.get("section_b") or {})
                _render_fmc_section_pdf(story, styles, _esc, _kv_table,
                                        "section_c", fmc.get("section_c") or {})
                _render_fmc_vitals_pdf(story, styles, _esc,
                                       fmc.get("section_d") or [])
                _render_fmc_treatment_pdf(story, styles, _esc, _kv_table,
                                          fmc.get("section_e") or {})
                _render_fmc_section_pdf(story, styles, _esc, _kv_table,
                                        "section_f", fmc.get("section_f") or {})
                _render_fmc_section_pdf(story, styles, _esc, _kv_table,
                                        "section_g", fmc.get("section_g") or {})
                story.append(Spacer(1, 6))

            vitals = p.get("vitals") or {}
            if any(vitals.values()):
                story.append(Paragraph("Vitalwerte", styles["SAFIRH2"]))
                vital_labels = {
                    "pulse": "Puls (bpm)",
                    "bp": "Blutdruck",
                    "resp_rate": "Atemfrequenz",
                    "spo2": "SpO2 (%)",
                    "temp": "Temperatur (°C)",
                    "gcs": "GCS",
                }
                vrows = []
                for key, label in vital_labels.items():
                    if vitals.get(key):
                        vrows.append([label, _esc(vitals[key])])
                if vrows:
                    story.append(_kv_table(vrows))
                story.append(Spacer(1, 6))

            injuries = p.get("injuries") or []
            if injuries:
                story.append(Paragraph("Verletzungen", styles["SAFIRH2"]))
                for inj in injuries:
                    story.append(Paragraph(f"• {_esc(inj)}", styles["Normal"]))
                story.append(Spacer(1, 6))

            treatments = p.get("treatments") or []
            medications = p.get("medications") or []
            if treatments or medications:
                story.append(Paragraph("Behandlungen / Medikamente", styles["SAFIRH2"]))
                for item in treatments:
                    val = item if isinstance(item, str) else (
                        item.get("description") or json.dumps(item, ensure_ascii=False)
                    )
                    story.append(Paragraph(f"• {_esc(val)}", styles["Normal"]))
                for item in medications:
                    val = item if isinstance(item, str) else (
                        item.get("name") or json.dumps(item, ensure_ascii=False)
                    )
                    story.append(Paragraph(f"• {_esc(val)}", styles["Normal"]))
                story.append(Spacer(1, 6))

            transcripts = p.get("transcripts") or []
            if transcripts:
                story.append(Paragraph("Transkripte", styles["SAFIRH2"]))
                for tr in transcripts:
                    if isinstance(tr, dict):
                        time_str = tr.get("time", "")
                        text = tr.get("text", "")
                    else:
                        time_str = ""
                        text = str(tr)
                    pf = (f"<b>[{_esc(time_str)}]</b> " if time_str else "") + _esc(text)
                    story.append(Paragraph(pf, styles["Normal"]))
                story.append(Spacer(1, 6))

            timeline = p.get("timeline") or []
            if timeline:
                story.append(Paragraph("Timeline", styles["SAFIRH2"]))
                for ev in timeline[-10:]:
                    if isinstance(ev, dict):
                        line = (
                            f"{_esc(ev.get('time',''))}  ·  [{_esc(ev.get('role',''))}]  "
                            f"{_esc(ev.get('event',''))}  —  {_esc(ev.get('details',''))}"
                        )
                    else:
                        line = _esc(ev)
                    story.append(Paragraph(line, styles["Normal"]))

    output_dir.mkdir(exist_ok=True, parents=True)
    filename = f"safir-patients-{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    filepath = output_dir / filename
    doc_pdf = SimpleDocTemplate(
        str(filepath),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title="SAFIR Patientendatenbank",
        author="SAFIR / CGI Deutschland",
    )
    doc_pdf.build(story)
    return filepath
