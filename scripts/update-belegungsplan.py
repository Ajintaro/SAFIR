#!/usr/bin/env python3
"""Aktualisiert das SAFIR GPIO Belegungsplan Word-Dokument mit der finalen
Verkabelung nach dem Hardware-Debug am 2026-04-09.

Was sich geändert hat:
- Beide Taster jetzt am 40-Pin Header (Power-Taster nicht mehr am J12)
- Konvention: weiß=GND, gelb=GPIO Input (vorher umgekehrt)
- LEDs werden über BC547 NPN Low-Side switches angesteuert (Pin 13/15)
- Externe 150Ω Pull-Up Widerstände an Pin 17 (3.3V) sind PFLICHT
- Pin 4 (5V) und Pin 6 (GND) für die LED-Schaltungen geteilt
"""
import sys
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor

DOC_PATH = Path("/home/jetson/cgi-afcea-san/docs/SAFIR-GPIO-Belegungsplan.docx")

doc = Document(DOC_PATH)


def set_table_cells(table, rows_data):
    """Schreibt mehrere Zeilen in eine bestehende Tabelle, fügt fehlende Rows hinzu."""
    while len(table.rows) < len(rows_data):
        table.add_row()
    while len(table.rows) > len(rows_data):
        table._tbl.remove(table.rows[-1]._tr)
    for r_idx, row_data in enumerate(rows_data):
        cells = table.rows[r_idx].cells
        # Falls Spaltenanzahl nicht passt, neue Spalten kann python-docx nicht
        # rückwirkend erweitern — wir gehen davon aus dass Spaltenanzahl stimmt
        for c_idx, val in enumerate(row_data):
            if c_idx < len(cells):
                cells[c_idx].text = str(val)


# ========================================================================
# Tabelle 1 — Status-Übersicht
# ========================================================================
table1 = doc.tables[0]
set_table_cells(table1, [
    ["Gerät", "Schnittstelle", "Status"],
    ["RC522 RFID-Reader (MFRC522 v2.0)", "Bit-Bang SPI (GPIO)", "Funktionsfähig"],
    ["SSD1306 OLED 128×64", "I2C Bus 7 (0x3C)", "Funktionsfähig"],
    ["Metzler K19-TF-W Taster A", "GPIO Pin 11 + LED via BC547 Pin 15", "Funktionsfähig"],
    ["Metzler K19-TF-W Taster B", "GPIO Pin 26 + LED via BC547 Pin 13", "Funktionsfähig"],
])

# ========================================================================
# Tabelle 2 — 40-Pin Header komplette Belegung
# ========================================================================
table2 = doc.tables[1]
set_table_cells(table2, [
    ["Pin", "Funktion", "Belegt durch", "Pin", "Funktion", "Belegt durch"],
    ["1",  "3.3V",          "OLED VCC (blau)",                   "2",  "5V",        "— frei —"],
    ["3",  "I2C8_SDA",      "OLED SDA (gelb)",                   "4",  "5V",        "BC547 LED-Versorgung (rot, geteilt)"],
    ["5",  "I2C8_SCL",      "OLED SCL (lila)",                   "6",  "GND",       "BC547 Emitter beide LEDs (schwarz)"],
    ["7",  "GPIO (RST)",    "RC522 RST (rot)",                   "8",  "UART1_TX",  "— frei —"],
    ["9",  "GND",           "OLED GND (grün)",                   "10", "UART1_RX",  "— frei —"],
    ["11", "GPIO Input",    "Taster A Schalter (gelb) + 150Ω→17","12", "GPIO",      "— frei —"],
    ["13", "GPIO Output",   "BC547 Basis Taster B LED (grün)",   "14", "GND",       "Taster A GND (weiß)"],
    ["15", "GPIO Output",   "BC547 Basis Taster A LED (blau)",   "16", "GPIO",      "— frei —"],
    ["17", "3.3V",          "RC522 VCC (braun) + 2× Pull-Up 150Ω","18","GPIO",      "— frei —"],
    ["19", "GPIO (MOSI)",   "RC522 MOSI (blau)",                 "20", "GND",       "RC522 GND (grau)"],
    ["21", "GPIO (MISO)",   "RC522 MISO (weiß)",                 "22", "GPIO (SCK)","RC522 SCK (schwarz)"],
    ["23", "SPI3 SCK",      "— frei —",                          "24", "GPIO (CS)", "RC522 NSS (lila)"],
    ["25", "GND",           "Taster B GND (weiß)",               "26", "GPIO Input","Taster B Schalter (gelb) + 150Ω→17"],
    ["27", "I2C0_SDA",      "EEPROM (reserviert)",               "28", "I2C0_SCL",  "EEPROM (reserviert)"],
    ["29", "GPIO",          "— frei —",                          "30", "GND",       "— frei —"],
    ["31", "GPIO",          "— frei —",                          "32", "GPIO",      "— frei —"],
    ["33", "GPIO",          "— frei —",                          "34", "GND",       "— frei —"],
    ["35", "GPIO",          "— frei —",                          "36", "GPIO",      "— frei —"],
    ["37", "GPIO",          "— frei —",                          "38", "GPIO",      "— frei —"],
    ["39", "GND",           "— frei —",                          "40", "GPIO",      "— frei —"],
])

# ========================================================================
# Tabelle 3 — RC522 Pinbelegung (Pin 22 statt 23 für SCK gefixt)
# ========================================================================
table3 = doc.tables[2]
set_table_cells(table3, [
    ["RC522 Pin", "Kabelfarbe", "Jetson Pin", "Funktion"],
    ["SDA (NSS)", "lila",       "Pin 24",     "GPIO CS (Bit-Bang)"],
    ["SCK",       "schwarz",    "Pin 22",     "GPIO SCK (Bit-Bang)"],
    ["MOSI",      "blau",       "Pin 19",     "GPIO MOSI (Bit-Bang)"],
    ["MISO",      "weiß",       "Pin 21",     "GPIO MISO (Bit-Bang)"],
    ["GND",       "grau",       "Pin 20",     "Ground"],
    ["RST",       "rot",        "Pin 7",      "GPIO Reset"],
    ["VCC",       "braun",      "Pin 17",     "3.3V (geteilt mit 2× Pull-Up)"],
])

# ========================================================================
# Tabelle 6 — OLED Pinbelegung (unverändert, bestätigt)
# ========================================================================
table6 = doc.tables[5]
set_table_cells(table6, [
    ["OLED Pin", "Kabelfarbe", "Jetson Pin", "Funktion"],
    ["SDA",      "gelb",       "Pin 3",      "I2C8_SDA"],
    ["SCL",      "lila",       "Pin 5",      "I2C8_SCL"],
    ["VCC",      "blau",       "Pin 1",      "3.3V"],
    ["GND",      "grün",       "Pin 9",      "Ground"],
])

# ========================================================================
# Tabelle 8 — Taster Kabelfarben (neue Konvention!)
# ========================================================================
table8 = doc.tables[7]
set_table_cells(table8, [
    ["Kabelfarbe", "Funktion"],
    ["weiß",       "Taster NO Kontakt 1 → GND"],
    ["gelb",       "Taster NO Kontakt 2 → GPIO Input (mit 150Ω Pull-Up)"],
    ["rot",        "LED + (Anode) → 5V"],
    ["schwarz",    "LED − (Kathode) → BC547 Kollektor"],
])

# ========================================================================
# Tabelle 9 — Taster A (war: Power J12, jetzt: Pin 11 mit LED via Pin 15)
# ========================================================================
table9 = doc.tables[8]
set_table_cells(table9, [
    ["Kabel",   "Farbe",   "Anschluss"],
    ["Schalter","weiß",    "Pin 14 (GND)"],
    ["Schalter","gelb",    "Pin 11 (GPIO Input, 150Ω Pull-Up zu Pin 17)"],
    ["LED +",   "rot",     "Pin 4 (5V, geteilt)"],
    ["LED −",   "schwarz", "BC547 Kollektor (Emitter → Pin 6 GND)"],
    ["Steuerleitung","blau","Pin 15 → BC547 Basis (HIGH = LED an)"],
])

# ========================================================================
# Tabelle 10 — Taster B (war: OLED-Nav Pin 11, jetzt: Pin 26 mit LED via Pin 13)
# ========================================================================
table10 = doc.tables[9]
set_table_cells(table10, [
    ["Kabel",   "Farbe",   "Anschluss"],
    ["Schalter","weiß",    "Pin 25 (GND)"],
    ["Schalter","gelb",    "Pin 26 (GPIO Input, 150Ω Pull-Up zu Pin 17)"],
    ["LED +",   "rot",     "Pin 4 (5V, geteilt)"],
    ["LED −",   "schwarz", "BC547 Kollektor (Emitter → Pin 6 GND)"],
    ["Steuerleitung","grün","Pin 13 → BC547 Basis (HIGH = LED an)"],
])

# ========================================================================
# Datum im Header / Stand aktualisieren
# ========================================================================
for p in doc.paragraphs:
    if "Stand:" in p.text:
        for run in p.runs:
            if "Stand:" in run.text or "April 2026" in run.text:
                run.text = "Stand: 09. April 2026"
        # Falls in einem einzelnen Run das ganze Datum steht
        if "Stand:" in p.text and len(p.runs) == 1:
            p.runs[0].text = "Stand: 09. April 2026"
        break

doc.save(DOC_PATH)
print(f"OK — {DOC_PATH.name} aktualisiert")
print(f"Tabellen geändert: 1, 2, 3, 6, 8, 9, 10")
